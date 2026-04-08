import streamlit as st
import pandas as pd
from openai import OpenAI
import time
import requests
from requests.auth import HTTPBasicAuth
import json
import re
import concurrent.futures
import urllib.parse
from tenacity import retry, stop_after_attempt, wait_exponential
from pydantic import BaseModel, Field, ValidationError, field_validator

# ==========================================
# 1. CONFIGURAÇÃO DA PÁGINA
# ==========================================
st.set_page_config(page_title="Arco Martech | Motor GEO", page_icon="🚀", layout="wide", initial_sidebar_state="collapsed")

# Lógica de Navegação via Query Parameters (Mais estável que botões)
query_params = st.query_params
if 'current_page' not in st.session_state:
    st.session_state['current_page'] = query_params.get("page", "Gerador de Artigos")
if 'show_inputs' not in st.session_state:
    st.session_state['show_inputs'] = False

# ==========================================
# ESTILOS GLOBAIS
# ==========================================
st.markdown("""
    <style>
    /* Importando as fontes do site da Arco */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Montserrat:wght@400;600;700;800&display=swap');

    /* Forçando a tipografia global */
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }

    h1, h2, h3 {
        font-family: 'Montserrat', sans-serif !important;
        font-weight: 700 !important;
        color: #111827 !important;
        letter-spacing: -0.02em;
    }

    /* ESCONDER COMPONENTES NATIVOS DO STREAMLIT */
    [data-testid="stSidebar"], header[data-testid="stHeader"] { display: none !important; }
    .block-container { padding-top: 1rem; max-width: 1200px; }

    .arco-tag {
        display: inline-flex;
        align-items: center;
        background-color: #E8F2FA;
        color: #418EDE !important;
        font-family: 'Montserrat', sans-serif;
        font-weight: 700;
        font-size: 0.75rem;
        letter-spacing: 0.05em;
        padding: 6px 16px;
        border-radius: 50px;
        text-transform: uppercase;
        margin-bottom: 1rem;
    }

    /* === 1. MENU PRINCIPAL COM LOGO ALINHADA === */
    /* Target estrito ao PRIMEIRO grupo de abas da página (O Menu). As sub-abas ignoram isso. */
    div[data-testid="stTabs"]:first-of-type > div > div[data-baseweb="tab-list"] {
        gap: 24px;
        border-bottom: 2px solid #E5E7EB;
        padding-left: 170px; /* Cria o espaço exato para a Logo */
        position: relative;
    }
    /* Injeta a Logo da Arco diretamente dentro da barra de abas principal */
    div[data-testid="stTabs"]:first-of-type > div > div[data-baseweb="tab-list"]::before {
        content: "";
        background-image: url('https://cdn.prod.website-files.com/6810e8cd1c64e82623876ba8/681134835142ef28e05b06ba_logo-arco-dark.svg');
        background-size: contain;
        background-repeat: no-repeat;
        background-position: left center;
        position: absolute;
        left: 0;
        top: 50%;
        transform: translateY(-50%);
        width: 140px;
        height: 35px;
    }
    div[data-testid="stTabs"]:first-of-type > div > div[data-baseweb="tab"] {
        font-family: 'Montserrat', sans-serif;
        font-weight: 600;
        color: #6B7280;
        padding-top: 16px;
        padding-bottom: 16px;
        background: transparent !important;
        border: none !important;
        box-shadow: none !important;
    }
    div[data-testid="stTabs"]:first-of-type > div > div[data-baseweb="tab"][aria-selected="true"] {
        color: #111827 !important;
        border-bottom: 3px solid #F05D23 !important; /* Laranja Arco */
        background: transparent !important;
    }

    /* === 2. TODOS OS BOTÕES PRIMÁRIOS (Quadrados normais 8px) === */
    div[data-testid="stButton"] > button[kind="primary"] {
        background-color: #111827 !important;
        color: #FFFFFF !important;
        border-radius: 8px !important; /* Retorna para o quadrado com canto leve */
        border: none !important;
        padding: 10px 24px !important;
        font-family: 'Inter', sans-serif;
        font-weight: 600 !important;
        height: 3.2em;
        transition: all 0.2s ease-in-out !important;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
    }
    div[data-testid="stButton"] > button[kind="primary"]:hover {
        background-color: #374151 !important;
        transform: translateY(-2px) !important;
    }
    div[data-testid="stButton"] > button[kind="primary"] *,
    div[data-testid="stButton"] > button[kind="primary"] p {
        color: #FFFFFF !important;
        fill: #FFFFFF !important;
        -webkit-text-stroke: 0px transparent !important;
        text-shadow: none !important;
    }

    /* === 3. BOTÃO HERÓI CIRCULAR DA HOME (Exclusivo) === */
    div[data-testid="stElementContainer"]:has(.hero-btn-hook) + div[data-testid="stElementContainer"] div[data-testid="stButton"] > button[kind="primary"] {
        border-radius: 50px !important; /* APENAS se tiver essa classe invisível ele fica redondo */
        height: 54px !important;
    }

    /* ESTILO DOS CARDS DE VENDA */
    .saas-card {
        background: #FFFFFF;
        border: 1px solid #E5E7EB;
        border-radius: 16px;
        padding: 24px;
        height: 100%;
        transition: transform 0.2s, box-shadow 0.2s;
    }
    .saas-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.1);
        border-color: #D1D5DB;
    }
    .card-title {
        font-family: 'Montserrat', sans-serif;
        font-weight: 700;
        font-size: 1.1rem;
        color: #111827;
        margin-bottom: 8px;
    }
    .card-text {
        font-size: 0.9rem;
        color: #4B5563;
        line-height: 1.5;
    }

    /* PIPELINE STYLING */
    .pipeline-container {
        font-family: 'Inter', sans-serif;
        font-size: 0.85em; 
        color: #6B7280; 
        text-align: center;
        margin: 2rem auto;
        background-color: #F9FAFB;
        padding: 12px;
        border-radius: 50px;
        border: 1px solid #E5E7EB;
        width: fit-content;
    }
    .pipeline-step {
        cursor: help; 
        color: #374151;
        font-weight: 500;
        transition: color 0.2s;
    }
    .pipeline-step:hover { color: #F05D23; }

    /* CONTAINER FLUTUANTE ÚNICO (LADO ESQUERDO) */
    .floating-controls-container {
        position: fixed;
        top: 110px; /* Ajuste para descer ou subir no eixo Y */
        left: 25px;
        z-index: 99999;
        display: flex;
        gap: 10px;
    }

    /* Botão de Ajuda (Vermelho) */
    div[data-testid="stPopover"]:first-child > button {
        background-color: #E21B22 !important;
        color: white !important;
        border-radius: 12px !important; /* Estilo mais moderno/quadrado suave */
        width: 50px !important;
        height: 50px !important;
        border: none !important;
        box-shadow: 0 4px 12px rgba(226, 27, 34, 0.3) !important;
    }

    /* Botão de Pautas (Laranja) - Usando seletor de irmão */
    div[data-testid="stPopover"]:nth-child(2) > button {
        background-color: #F05D23 !important;
        color: white !important;
        border-radius: 12px !important;
        width: 50px !important;
        height: 50px !important;
        border: none !important;
        box-shadow: 0 4px 12px rgba(240, 93, 35, 0.3) !important;
    }

    div[data-testid="stPopover"] button p {
        font-size: 22px !important;
        font-weight: bold;
        margin: 0 !important;
    }
    /* Remove as bordas e fundos dos botões do menu */
    div[data-testid="stButton"] > button[kind="secondary"] {
        border: none !important;
        background: transparent !important;
        box-shadow: none !important;
    }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# 1.1 MENU DE NAVEGAÇÃO SAAS NO TOPO
# ==========================================
nav_cols = st.columns([2, 2, 2, 2, 2, 2])

with nav_cols[0]:
    st.markdown('<img src="https://cdn.prod.website-files.com/6810e8cd1c64e82623876ba8/681134835142ef28e05b06ba_logo-arco-dark.svg" style="width: 140px; margin-top: -20px;" alt="Logo Arco">', unsafe_allow_html=True)

opcoes_menu = ["Gerador de Artigos", "BrandBook", "Monitor de GEO", "Revisor de GEO", "Auditor de Artigos"]

# Aplicamos o estilo do menu selecionado DE UMA VEZ AQUI EM CIMA, para não empurrar os botões no loop
try:
    index_selecionado = opcoes_menu.index(st.session_state['current_page'])
    # Usa stroke pra bold sem mudar a largura, e border-bottom na cor laranja
    st.markdown(f"""
    <style>
    div[data-testid="stHorizontalBlock"]:first-of-type div[data-testid="stColumn"]:nth-child({index_selecionado + 2}) button {{
        color: #111827 !important; 
        -webkit-text-stroke: 0.6px #111827 !important;
        border-bottom-color: #F05D23 !important;
    }}
    </style>
    """, unsafe_allow_html=True)
except ValueError:
    pass

# Agora o loop só renderiza os botões, sem injetar tags extras no meio do caminho
for i, opcao in enumerate(opcoes_menu):
    with nav_cols[i+1]:
        if st.button(opcao, use_container_width=True, key=f"nav_{i}"):
            st.session_state['current_page'] = opcao
            st.rerun()

st.markdown("<div style='margin-bottom: 0.5rem;'></div>", unsafe_allow_html=True)

# ==========================================
# BOTÃO FLUTUANTE DE AJUDA (ESQUERDA)
# ==========================================
st.markdown('<div class="floating-controls-container">', unsafe_allow_html=True)
with st.popover("?"):
    st.header("📖 Guia Prático do Motor")
    st.markdown("Bem-vindo à v7.0. Este motor funciona como sua **equipe particular de especialistas**. Ele espiona a concorrência, entende as regras do Google e das IAs, e escreve conteúdos usando a voz exata da sua marca.")
    
    with st.expander("🚀 Como usar as 5 Abas?"):
        st.markdown("""
        **1. Gerador:** Cria artigos completos do zero. Você dá o tema (e links de referência se quiser), ele pesquisa o mercado e redige.
        
        **2. Brandbook:** O 'cérebro' do sistema. É aqui que dizemos o que cada marca da Arco pode ou não falar.
        
        **3. Monitor:** Ferramenta de auditoria. Cole um texto qualquer aqui para a IA dar uma nota de confiabilidade e sugerir melhorias.
        
        **4. Adaptador & Revisor:** Transforme seus E-books/PDFs em artigos "Teaser" para captar Leads, ou conserte textos antigos do blog para voltarem a ranquear.
        
        **5. Auditor de Visibilidade:** Coloque o link de um artigo seu e descubra se o Google ou as IAs já estão recomendando ele.
        """)
        
    with st.expander("📚 O que significam as Notas Matemáticas?"):
        st.markdown("""
        O nosso motor avalia seu texto em duas frentes: **Estrutura** e **Autoridade**.
        
        **Notas de Estrutura:**
        * **Chunk Citability (Legibilidade):** Mede se o texto é fácil de ler. Parágrafos curtos, listas e frases de impacto aumentam a nota.
        * **Answer-First:** Avalia se você enrolou ou se entregou a resposta principal logo no começo do texto.
        
        **Notas de Autoridade:**
        * **Evidence Density (Evidências):** Mede se você usou números, estatísticas reais e links para provar o que diz.
        * **Information Gain (Ineditismo):** Calcula o quanto de informação nova você trouxe em relação ao que já existe no Top 3 do Google.
        * **Entity Coverage:** Avalia se você usou o vocabulário que todo especialista do seu nicho deveria usar.
        """)
        
    with st.expander("🤖 O que são os Testes de IA?"):
        st.markdown("""
        Nós simulamos como o ChatGPT ou Perplexity julgariam o seu texto:
        
        * **Retrieval Simulation:** É a chance de uma IA escolher o seu texto como fonte oficial para responder a um usuário.
        * **Risco de Hijacking:** Mede o risco de um concorrente "roubar" o seu clique por ter explicado o assunto de forma mais direta e didática que você.
        """)

@st.cache_data(ttl=3600, show_spinner=False)
def buscar_trending_topics_educacao():
    """Busca pautas em múltiplas fontes simultaneamente de forma rápida"""
    fontes_rss = [
        "https://news.google.com/rss/search?q=MEC+OR+ENEM+OR+escolas&hl=pt-BR&gl=BR&ceid=BR:pt-419", 
        "https://news.google.com/rss/search?q=edtech+OR+gestão+escolar+OR+inadimplência+escolar&hl=pt-BR&gl=BR&ceid=BR:pt-419",
        "https://g1.globo.com/rss/g1/educacao/" 
    ]
    
    def extrair_noticia(url):
        try:
            import feedparser
            feed = feedparser.parse(url)
            if feed.entries:
                entry = feed.entries[0]
                titulo_limpo = entry.title.split(' - ')[0].strip()
                titulo_curto = titulo_limpo[:55] + "..." if len(titulo_limpo) > 55 else titulo_limpo
                # Agora retorna uma tupla: (Título, Link)
                return (titulo_curto, entry.link)
        except Exception:
            return None
        return None

    with concurrent.futures.ThreadPoolExecutor() as executor:
        resultados = list(executor.map(extrair_noticia, fontes_rss))
        
    # Usamos set() para remover duplicatas baseadas na tupla inteira
    pautas_coletadas = list(set([res for res in resultados if res]))
    
    # Adicionamos um link vazio nas de fallback
    pautas_fallback = [
        ("Inovação tecnológica na gestão escolar", ""), 
        ("Uso de IA no dia a dia da sala de aula", ""), 
        ("Estratégias para retenção de alunos", "")
    ]
    
    return (pautas_coletadas + pautas_fallback)[:3]
        
# 2. O botão de Pautas (🔥) agora SÓ aparece se o formulário estiver aberto
if st.session_state.get('show_inputs', False) and st.session_state.get('current_page') == "Gerador de Artigos":
    with st.popover("🔥"):
        st.markdown("### 🔥 Pautas em Alta")
        st.caption("Tendências detectadas via Google News e MEC agora:")
        
        # Chama a função que agora retorna (Titulo, Link)
        pautas_quentes = buscar_trending_topics_educacao()
        
        for pauta, link in pautas_quentes:
            col_btn, col_link = st.columns([8, 2])
            
            with col_btn:
                # O botão continua preenchendo o input lá embaixo
                if st.button(f"{pauta}", use_container_width=True, key=f"trend_{pauta}"):
                    st.session_state['pauta_sugerida'] = pauta
                    st.rerun()
                    
            with col_link:
                # Se houver um link real, mostra o botão de abrir aba
                if link:
                    st.markdown(
                        f"""<a href="{link}" target="_blank" title="Ler notícia original" 
                        style="display: flex; align-items: center; justify-content: center; 
                        height: 100%; text-decoration: none; font-size: 1.2rem; 
                        background-color: #F3F4F6; border-radius: 8px;">🔗</a>""", 
                        unsafe_allow_html=True
                    )


# Armazenando o HTML do pipeline para usar depois
pipeline_html = """
<div class="pipeline-container">
    <strong style="color: #111827; font-family: 'Montserrat', sans-serif;">O Caminho do Conteúdo:</strong> 
    <span title="1. Pesquisa: Espiona o Top 3 do Google e as IAs (como ChatGPT) já dizem sobre o tema." class="pipeline-step">1. Pesquisa</span> ➔ 
    <span title="2. Intenção: Descobre a verdadeira dúvida por trás das buscas (o que o leitor quer saber)." class="pipeline-step">2. Intenção</span> ➔ 
    <span title="3. Vocabulário: Mapeia os jargões e conceitos obrigatórios de autoridade." class="pipeline-step">3. Vocabulário</span> ➔ 
    <span title="4. Escrita: Redige o texto usando o tom de voz e regras anti-IA." class="pipeline-step">4. Escrita</span> ➔ 
    <span title="5. Código SEO: Cria os dados ocultos (Schema) para o Google." class="pipeline-step">5. Código SEO</span> ➔ 
    <span title="6. Auditoria: Calcula notas de leitura, resposta direta e evidências." class="pipeline-step">6. Auditoria</span> ➔ 
    <span title="7. Teste de IAs: Simula se o seu texto está bom para virar fonte do SGE." class="pipeline-step">7. Teste de IAs</span>
</div>
"""

# ==========================================
# ESTRUTURAS PYDANTIC
# ==========================================
class MetadadosArtigo(BaseModel):
    title: str = Field(..., description="Título H1 otimizado (max 60 chars)")
    meta_description: str = Field(..., description="Meta description persuasiva (max 150 chars)")
    dicas_imagens: list[str] = Field(..., description="Lista com 2 prompts curtos EM INGLÊS para buscar imagens (ex: ['futuristic classroom', 'student studying'])")
    schema_faq: dict = Field(..., description="Objeto JSON-LD FAQPage completo e idêntico ao texto")

    @field_validator('title', mode='before')
    @classmethod
    def ajustar_tamanho_titulo(cls, v: str) -> str:
        # NOVA REGRA: Arranca anos (ex: 2024, 2025, 2026, etc) do título
        v_limpo = re.sub(r'\b202[4-9]\b', '', v)
        # Limpa hifens, dois pontos ou espaços que sobrarem soltos no final
        v_limpo = re.sub(r'[-\s:]+$', '', v_limpo).strip()
        return v_limpo
        
    @field_validator('meta_description', mode='before')
    @classmethod
    def ajustar_tamanho_meta(cls, v: str) -> str:
        return v[:147] + "..." if len(v) > 150 else v


# ==========================================
# 2. BRANDBOOK EMBUTIDO 
# ==========================================
if 'brandbook_df' not in st.session_state:
    dados_iniciais = [
        {
            "Marca": "SAS Educação",
            "URL": "https://www.saseducacao.com.br/",
            "Posicionamento": "Marca visionária, líder em aprovação. Entrega de valor em tecnologia e serviço. | Protagonistas na evolução da forma de ensinar e aprender. Abordagem com diagnósticos e embasamentos profundos, superamos as expectativas de parceiros. Somos alta performance e transformamos complexidade em oportunidades. | Promessa: Educação de excelência com foco em resultados acadêmicos, suporte pedagógico próximo e uso de dados para aprendizado.",
            "Territorios": "Vestibulares, Tecnologia, Inovação, Pesquisas",
            "TomDeVoz": "Acadêmico, inovador, especialista e inspirador. Visionário, colaborativo",
            "PublicoAlvo": "Mantenedores e gestores de escolas médias e grandes, com alto rigor acadêmico e foco em resultados no ENEM. Estudantes, vestibulandos e pais",
            "RegrasNegativas": "Não usar tom professoral antiquado, não prometer aprovação sem esforço.",
            "RegrasPositivas": "Destaque os diferenciais: - Líder nacional em aprovação no SiSU 2025, - Maior sistema de ensino do Brasil, - +1.300 escolas parceiras, - 97% de fidelização. Propósito da marca: Moldar, com coragem e embasamento, a educação do futuro ao lado das escolas."
        },
        {
            "Marca": "Geekie",
            "URL": "https://www.geekie.com.br/",
            "Posicionamento": "Metodologia inovadora (aluno no centro), fácil de implementar. | Material didático inteligente que apoia práticas ativas e que possibilita a personalização da aprendizagem por meio de dados. | Promessa: Aprendizado personalizado, engajante e baseado em dados.",
            "Territorios": "Inovação, IA/Personalização, Tecnologia, Dados",
            "TomDeVoz": "Inovador, moderno, ágil. Transformador, visionário, experimental, adaptável e inspirador",
            "PublicoAlvo": "Mantenedores e Gestores. Diretores de inovação e escolas modernas.",
            "RegrasNegativas": "Não parecer sistema engessado, não usar linguagem punitiva.",
            "RegrasPositivas": "Destaque os diferenciais: - A primeira plataforma de educação baseada em dados, - Mais de 12 milhões de estudantes impactados, - Melhor solução de IA premiada no Top Educação. Propósito da marca: Transformar a educação para que cada estudante seja tratado como único."
        },
        {
            "Marca": "COC",
            "URL": "https://coc.com.br/",
            "Posicionamento": "Marca aprovadora que evolui a escola pedagogicamente. | Promover transformação de alto impacto, através de resultados de crescimento para a gestão da escola e ao longo de toda a trajetória do aluno | Promessa: Resultados de crescimento para a gestão da escola e ao longo de toda a trajetória do aluno.",
            "Territorios": "Vestibulares, Esportes, Gestão escolar, Crescimento",
            "TomDeVoz": "Consultivo, parceiro, dinâmico. Viva, ponta firme, sagaz, aberta, contemporânea",
            "PublicoAlvo": "Mantenedores e Gestores. Coordenadores pedagógicos.",
            "RegrasNegativas": "Jamais foque o discurso em 'desempenho agregado de escola', o foco principal do benefício deve ser sempre o Aluno. Nunca repita exaustivamente 'O Laboratório de Redação COC', abrevie ou omita o 'COC' após a primeira menção. Nunca use 'plataforma educacional', prefira chamá-la de 'ferramenta'.",
            "RegrasPositivas": "Destaque os diferenciais: - Mais de 60 anos, - Melhor consultoria do Brasil 2x premiada no Top Educação. Propósito: Impulsionar escolas rumo a uma educação contemporânea de excelência."
        },
        {
            "Marca": "Sistema Positivo",
            "URL": "https://www.sistemapositivo.com.br/",
            "Posicionamento": "Formação integral, humana e próxima. A maior rede do Brasil. | Com uma abordagem inspiradora e humana, somos referência em solutions que guiam nossas escolas parceiras a evoluírem na missão de ensinar, transformando positivamente a vida dos brasileiros.",
            "Territorios": "Formação integral, Inclusão, Tradição",
            "TomDeVoz": "Acolhedor, tradicional, humano. Experiente, criativa, inovadora e segura",
            "PublicoAlvo": "Famílias. Mantenedores e Gestores de escolas tradicionais.",
            "RegrasNegativas": "Não parecer frio, não usar jargões técnicos sem contexto acolhedor.",
            "RegrasPositivas": "Destaque os diferenciais: - Mais de 45 anos de atuação. Propósito: Inspirar e fortalecer escolas para que evoluam a educação brasileira com humanidade."
        },
        {
            "Marca": "SAE Digital",
            "URL": "https://sae.digital/",
            "Posicionamento": "Melhor integração físico/digital, hiperatualizada. | Nos consolidamos como o sistema de ensino atualizado, que melhor integra o físico com o digital para potencializar o resultado dos alunos e dos nossos parceiros.",
            "Territorios": "Tecnologia, Inovação Digital",
            "TomDeVoz": "Prático, tecnológico, dinâmico. Jovem, amigável, antenado, parceiro",
            "PublicoAlvo": "Mantenedores e Gestores buscando modernização com custo-benefício.",
            "RegrasNegativas": "Não parecer inacessível, não diminuir a importância do material físico.",
            "RegrasPositivas": "Propósito: Desbravar o caminho para uma educação excelente e acessível, que permita a cada aluno e educador escolher e concretizar seus sonhos."
        },
        {
            "Marca": "Conquista Solução Educacional",
            "URL": "https://www.educacaoconquista.com.br/",
            "Posicionamento": "Solução completa focada na parceria Escola-Família. | Desenvolvimento integral e acessível, a partir de 4 pilares: educação financeira, empreendedorismo, educação socioemocional e família.",
            "Territorios": "Família, Educação Infantil, Valores, Comunidade, Empreendedorismo, Socioemocional",
            "TomDeVoz": "Familiar, parceiro, simples e didático. Integradora, descomplicada",
            "PublicoAlvo": "Pais. Mantenedores e Gestores de escolas de educação infantil.",
            "RegrasNegativas": "Não usar tom corporativo frio, não focar em pressão de vestibular.",
            "RegrasPositivas": "Propósito: Colaborar com escolas para formar alunos protagonistas que constroem seu próprio caminho."
        },
        {
            "Marca": "Escola da Inteligência",
            "URL": "https://escoladainteligencia.com.br/",
            "Posicionamento": "Um ecossistema de educação que transforma alunos, professores, escolas e famílias pelo desenvolvimento da inteligência socioemocional.",
            "Territorios": "Comunidade, Socioemocional, habilidades e competências",
            "TomDeVoz": "Madura, especialista, profunda, humana, acessível, sentimental, suave, estável.",
            "PublicoAlvo": "Mantenedores e Gestores de escolas médias, tradicionais que desejam qualidade e são movidos por um senso de propósito (Ticket alto).",
            "RegrasNegativas": "Evitar linguagem robótica, sem focar excessivamente na competição e em pressões externas.",
            "RegrasPositivas": "Destaque: Primeira solução socioemocional do mercado Brasileiro, presente desde 2010. Tricampeões invictos do Top Educação. Citar ferramentas 'Pulso', 'Mapa Socioemocional' e 'Indicadores Multifocais'. 1.2 milhões de pessoas impactadas."
        },
        {
            "Marca": "PES English",
            "URL": "https://www.pesenglish.com.br/",
            "Posicionamento": "O maior programa de inglês integrado às escolas, facilitador do ensino de qualidade, com resultados que mudam vidas. | Promessa: Educação acessível, integrada e descomplicada.",
            "Territorios": "Bilíngue, crescimento, tecnologia",
            "TomDeVoz": "Especialista, humano, dinâmico, acessível, suave",
            "PublicoAlvo": "Mantenedores e Gestores de escolas que visam escala na educação linguística com custo-benefício para famílias.",
            "RegrasNegativas": "Não prometer fluência irreal em curto prazo, não utilizar termos em inglês soltos sem conexão com o currículo.",
            "RegrasPositivas": "Destaque: 91% de aprovação nos exames de Cambridge, parcerias com Cambridge e Pearson, sistema 'Level Up'. Programa curricular flexível. Mais de 800 escolas, custando 10x menos que curso de idiomas avulso."
        },
        {
            "Marca": "Nave a Vela",
            "URL": "https://www.naveavela.com.br/",
            "Posicionamento": "Referência em educação tecnológica para formar estudantes protagonistas na resolução de problemas reais com tecnologia e criatividade por meio de experiências práticas.",
            "Territorios": "Inovação, tecnologia, criatividade",
            "TomDeVoz": "Especialista, espontâneo, racional, dinâmico",
            "PublicoAlvo": "Mantenedores e Gestore de escolas modernas que valorizam cultura Maker e letramento tecnológico.",
            "RegrasNegativas": "Não desmerecer o ensino tradicional. O foco deve ser a integração complementar.",
            "RegrasPositivas": "Destaque: Abordagem STEAM, 4Cs (criatividade, pensamento crítico, colaboração e comunicação), foco em Inteligência Artificial ética. 4x ganhadores no Top Educação em Educação Tecnológica."
        },
        {
            "Marca": "Programa Pleno",
            "URL": "https://programapleno.com.br/",
            "Posicionamento": "O Pleno transforma o convívio escolar através da educação socioemocional interdisciplinar e com rigor científico, trabalhando saúde mental, física e relações interpessoais.",
            "Territorios": "Projetos, socioemocional, habilidades e competências, bem estar",
            "TomDeVoz": "Coletivo, jovem, dinâmico, espontâneo, sofisticado, humano, especialista",
            "PublicoAlvo": "Mantenedores e Gestores buscando metodologias baseadas em projetos com comprovação científica.",
            "RegrasNegativas": "Não atrelar as soluções como um serviço clínico. É um desenvolvimento escolar de convivência.",
            "RegrasPositivas": "Destaque: Baseado no modelo internacional CASEL, abordagem SAFER, aprendizado baseado em projetos, Guia de trabalho nos espaços públicos e alinhamento à BNCC."
        },
        {
            "Marca": "Gênio das Finanças",
            "URL": "https://geniodasfinancas.com.br/",
            "Posicionamento": "Através da educação financeira comportamental, unimos escolas, alunos e famílias para cultivar autonomia, consciência e equilíbrio nas decisões financeiras, fortalecendo projetos de vida mais saudáveis.",
            "Territorios": "Educação financeira comportamental, habilidades e competências",
            "TomDeVoz": "Dinâmico, specialist, acessível, humano, estável",
            "PublicoAlvo": "Mantenedores e Gestores de escolas focadas em habilidades para a vida do aluno do ensino básico.",
            "RegrasNegativas": "Não usar termos como ficar rico ou fórmulas mágicas. O foco é 'comportamental e equilíbrio', nunca promessas milagrosas.",
            "RegrasPositivas": "Destaque: Educação financeira com propósito, ensinando finanças sem julgamentos e com foco no bem-estar emocional."
        },
        {
            "Marca": "Maralto",
            "URL": "https://maralto.com.br/",
            "Posicionamento": "A Maralto assume a sua responsabilidade no processo de construção de um país leitor e apresenta o Programa de Formação Leitora Maralto com o desejo de promover diálogos em torno do livro, da leitura e dos leitores.",
            "Territorios": "Literatura, associação pedagógica",
            "TomDeVoz": "Coletiva, especialista, sofisticada, humana, profunda, formal",
            "PublicoAlvo": "Educadores que apreciam bibliotecas robustas e incentivo literário profundo.",
            "RegrasNegativas": "Não resumir a literatura a apenas materiais didáticos conteudistas. A chave é 'leitura por prazer e diálogo'.",
            "RegrasPositivas": "Destaque: Investimento autoral em conteúdo literário e visual. Propósito: Formar um país de leitores."
        },
        {
            "Marca": "International School",
            "URL": "https://internationalschool.global/",
            "Posicionamento": "O programa bilíngue mais premiado do Brasil. Pioneira em bilinguismo no país. Prover soluções educacionais consistentes e inovadoras. Transformar vidas por meio da educação bilíngue. Empoderar a comunidade escolar para desenvolver o aluno como ser integral. | Promessa: Resultados concretos no aprendizado.",
            "Territorios": "Bilinguismo, educação, integral, viagens, inovação, pioneirismo",
            "TomDeVoz": "Especialista, inovador, inspirador, prático, pioneiro, parceiro",
            "PublicoAlvo": "Gestores, diretores e coordenadores de escolas. Pais e famílias. Escolas privadas de ticket alto e famílias de classes A, B e C.",
            "RegrasNegativas": "Não usar termos genéricos sem contexto, não soar arrogante ou sabe-tudo. Não inferir que quem aprende inglês é superior ou melhor. Não citar palavras em inglês sem tradução entre parênteses depois. Não focar o discurso somente nos pais (lembrar sempre da figura da escola). NUNCA usar a construção 'neste artigo iremos' ou similares.",
            "RegrasPositivas": "Focar em estrutura informativa. Sempre trazer dados para embasar afirmações vindos de fontes seguras e confiáveis, sempre citar e linkar a fonte dos dados, preferir fontes de pesquisas, governos e instituições de renome. Sempre começar o primeiro parágrafo com um gancho que instigue a leitura, de preferência acompanhado de dado. Podemos usar pesquisas nacionais ou internacionais. Sempre usar construção gramatical focada em clareza: iniciar parágrafos com frases de afirmação, não com conectivos. Sempre conectar com a importância de aprender inglês indo além da gramática: focar na importância de aprender com contexto. Destaque os diferenciais (CSV): Utilização da metodologia CLIL de forma integral. Aborde vivências internacionais reais (KSCIA, Cambridge, Minecraft, Ubisoft, Leo) e a integração do inglês à rotina escolar."
        },
        {
            "Marca": "Isaac",
            "URL": "https://isaac.com.br/",
            "Posicionamento": "A maior plataforma financeira e de gestão para a educação. | Promessa: Mensalidades em dia, sem dor de cabeça.",
            "Territorios": "Gestão financeira, Inovação, dados, tecnologia",
            "TomDeVoz": "Corporativo, direto, analítico. Simples (acessível) e parceiro, especialista em gestão financeira.",
            "PublicoAlvo": "Mantenedores, gestores e diretores financeiros de escolas, faculdades e confessionais.",
            "RegrasNegativas": "Não parecer banco engessado, não usar linguagem infantilizada ou agressiva contra a família devedora.",
            "RegrasPositivas": "Destaque: Diminuição real da inadimplência, 2x premiada no Top educação, excelência técnica, comprometimento e resultados tangíveis."
        },
        {
            "Marca": "ClassApp",
            "URL": "https://www.classapp.com.br/",
            "Posicionamento": "A agenda escolar online melhor avaliada do Brasil | Promessa: Mais que funcionalidades, soluções definitivas para os desafios reais da escola.",
            "Territorios": "Comunicação escolar, gestão, inovação",
            "TomDeVoz": "Autoridade acessível (sabe e explica como faz), empática e humana.",
            "PublicoAlvo": "Mantenedores, gestores, diretores, coordenadores, TI e marketing de escolas.",
            "RegrasNegativas": "Não falar mal do uso do papel de forma grosseira, sempre usar como avanço de modernização.",
            "RegrasPositivas": "Destaque: Adesão de 95% e leitura de 85%, segurança, única vencedora do Top Educação na categoria e mais de 260 mil avaliações com nota 4.8."
        },
        {
            "Marca": "Activesoft",
            "URL": "https://activesoft.com.br/",
            "Posicionamento": "Gestão escolar mais simples e eficiente com a Activesoft: tudo o que sua escola precisa para otimizar processos, ganhar eficiência e alcançar melhores resultados.",
            "Territorios": "Gestão escolar, dados, gestão acadêmica, gestão financeira, administrativa",
            "TomDeVoz": "Simples, acessível, clara e amigável.",
            "PublicoAlvo": "Mantenedores, gestores, diretores e TI de escolas.",
            "RegrasNegativas": "Não usar terminologia muito rebuscada para TI.",
            "RegrasPositivas": "Destaque: Plataforma 100% online (ao contrário de desktops), 25 anos de mercado, atendimento em chat em até 2 minutos (90% de satisfação). Mais de 3 milhões de usuários."
        },
        {
            "Marca": "Arco Educação",
            "URL": "https://www.arcoeducacao.com.br/",
            "Posicionamento": "A plataforma integrada de soluções educacionais da Arco Educação. Ponto de encontro de soluções que simplificam a rotina. +12.000 escolas parceiras e +4 milhões de alunos. | Promessa: Tudo que a educação precisa, em um só lugar.",
            "Territorios": "Conexão e tecnologia, foco no elo entre gestão e família (herança isaac/ClassApp).",
            "TomDeVoz": "Confiável, estratégica: torna o complicado mais simples, conecta o que estava separado.",
            "PublicoAlvo": "Mantenedores, gestores e diretores. Professores. Famílias. Alunos.",
            "RegrasNegativas": "Não apresentar como um simples repositório, mas como um ecossistema.",
            "RegrasPositivas": "Destaque: Apenas uma marca com o tamanho e história da Arco conseguiria reunir o melhor de pedagógico, gestão e tecnologia em um só lugar."
        }
    ]
    st.session_state['brandbook_df'] = pd.DataFrame(dados_iniciais)

# ==========================================
# 2.1 BASE DE DADOS DOS ESPECIALISTAS (GHOSTWRITING)
# ==========================================
if 'especialistas_df' not in st.session_state:
    dados_especialistas = [
        {"Especialista": "Professor Idelfranio Moreira De Sousa", "Link do Artigo": "https://exemplo.com"},
        {"Especialista": "Professor Ademar Celedonio Guimaraes Junior", "Link do Artigo": "https://www.linkedin.com/pulse/alunos-mais-ricos-do-brasil-t%25C3%25AAm-notas-inferiores-aos-celed%25C3%25B4nio-g-jr-eav6f/"},
        {"Especialista": "Professor Ademar Celedonio Guimaraes Junior", "Link do Artigo": "https://www.linkedin.com/pulse/como-atualidades-podem-ser-cobradas-enem-ademar-celed%25C3%25B4nio-g-jr-cjedf/"},
        {"Especialista": "Professor Ademar Celedonio Guimaraes Junior", "Link do Artigo": "https://www.linkedin.com/pulse/como-educa%25C3%25A7%25C3%25A3o-do-futuro-pode-ser-moldada-partir-uso-celed%25C3%25B4nio-g-jr-4cl7f/"},
        {"Especialista": "Professor Ademar Celedonio Guimaraes Junior", "Link do Artigo": "https://www.linkedin.com/pulse/l%25C3%25ADderes-que-moldam-vidas-celebrando-o-dia-do-diretor-celed%25C3%25B4nio-g-jr-iyizf/"},
        {"Especialista": "Professor Ademar Celedonio Guimaraes Junior", "Link do Artigo": "https://www.linkedin.com/pulse/vacinas-de-mrna-da-rejei%25C3%25A7%25C3%25A3o-acad%25C3%25AAmica-ao-pr%25C3%25AAmio-nobel-ademar/"},
        {"Especialista": "Professor Ademar Celedonio Guimaraes Junior", "Link do Artigo": "https://www.linkedin.com/pulse/5-formas-de-investir-na-educa%25C3%25A7%25C3%25A3o-do-seu-filho-e-o-celed%25C3%25B4nio-g-jr/"},
        {"Especialista": "Professor Ademar Celedonio Guimaraes Junior", "Link do Artigo": "https://www.linkedin.com/pulse/construindo-repert%25C3%25B3rio-cultural-para-o-enem-e-fuvest-celed%25C3%25B4nio-g-jr/"},
        {"Especialista": "Professor Ademar Celedonio Guimaraes Junior", "Link do Artigo": "https://www.linkedin.com/pulse/bncc-e-educa%25C3%25A7%25C3%25A3o-midi%25C3%25A1tica-ferramentas-cruciais-em-um-celed%25C3%25B4nio-g-jr/"},
        {"Especialista": "Professor Ademar Celedonio Guimaraes Junior", "Link do Artigo": "https://www.linkedin.com/pulse/quanto-maior-o-investimento-em-tecnologia-ser%25C3%25A1-de-celed%25C3%25B4nio-g-jr/"},
        {"Especialista": "Professor Ademar Celedonio Guimaraes Junior", "Link do Artigo": "https://www.linkedin.com/pulse/censo-escolar-2025-brasil-perde-11-milh%C3%A3o-de-alunos-ademar-m1oae/"}
    ]
    st.session_state['especialistas_df'] = pd.DataFrame(dados_especialistas)
    
# ==========================================
# 3. CONEXÃO SEGURA E CREDENCIAIS
# ==========================================
try:
    TOKEN = st.secrets["OPENROUTER_KEY"]
except Exception:
    TOKEN = None

try:
    SERPAPI_KEY = st.secrets["SERPAPI_KEY"]
except Exception:
    SERPAPI_KEY = None

def obter_credenciais_cms(marca):
    """Busca as credenciais (WP ou Drupal) da marca nos secrets."""
    try:
        if "wordpress" in st.secrets and marca in st.secrets["wordpress"]:
            creds = st.secrets["wordpress"][marca]
            return creds.get("WP_URL", ""), creds.get("WP_USER", ""), creds.get("WP_APP_PASSWORD", ""), creds.get("CMS_TYPE", "wp").lower()
    except Exception:
        pass
    return "", "", "", "wp"
# ==========================================
# 3.2 FUNÇÕES DE CONTEXTO E BUSCA
# ==========================================
   
@st.cache_data(ttl=3600, show_spinner=False)
def buscar_contexto_google(palavra_chave):
    if not SERPAPI_KEY:
        return "Sem chave Serper configurada. Pule o contexto do Google."
    url = "https://google.serper.dev/search"
    payload = json.dumps({"q": palavra_chave, "gl": "br", "hl": "pt-br"})
    headers = {'X-API-KEY': SERPAPI_KEY, 'Content-Type': 'application/json'}
    try:
        response = requests.request("POST", url, headers=headers, data=payload)
        dados = response.json()
        contexto_extraido = []
        if "answerBox" in dados:
            snippet = dados["answerBox"].get("snippet") or dados["answerBox"].get("answer", "Sem texto")
            contexto_extraido.append(f"📍 GOOGLE FEATURED SNIPPET ATUAL:\n{snippet}\n")
        if "knowledgeGraph" in dados:
            desc = dados["knowledgeGraph"].get("description", "")
            contexto_extraido.append(f"🧠 GOOGLE KNOWLEDGE GRAPH:\n{desc}\n")
        if "organic" in dados:
            contexto_extraido.append("📊 TOP 3 RESULTADOS ORGÂNICOS (CONTEÚDO LIDO VIA JINA):")

            def buscar_jina(res_item, index):
                titulo = res_item.get('title', 'Sem Título')
                snippet = res_item.get('snippet', 'Sem Snippet')
                link = res_item.get('link', '')
                conteudo_real = ""
                if link:
                    try:
                        jina_headers = {
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                            'X-Return-Format': 'markdown',
                            'Accept': 'text/plain'
                        }
                        jina_res = requests.get(f"https://r.jina.ai/{link}", headers=jina_headers, timeout=12)
                        if jina_res.status_code == 200:
                            conteudo_real = jina_res.text[:1500]
                    except Exception:
                        conteudo_real = "Falha ao ler o conteúdo integral."
                return f"{index+1}. Título: {titulo}\n Snippet: {snippet}\n Link: {link}\n Conteúdo:\n{conteudo_real}\n"

            with concurrent.futures.ThreadPoolExecutor() as executor:
                resultados_jina = list(executor.map(lambda x: buscar_jina(x[1], x[0]), enumerate(dados["organic"][:3])))
            contexto_extraido.extend(resultados_jina)
        resultado_final = "\n".join(contexto_extraido)
        return resultado_final if resultado_final else "Sem resultados orgânicos relevantes."
    except Exception as e:
        return f"Erro ao coletar dados do Google (Serper): {e}"

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def chamar_llm(system_prompt, user_prompt, model, temperature=0.3, response_format=None):
    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=TOKEN,
        default_headers={"HTTP-Referer": "https://arcomartech.com", "X-Title": "Gerador GEO WP"}
    )
    kwargs = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": temperature,
    }
    if response_format:
        kwargs["response_format"] = response_format
    response = client.chat.completions.create(**kwargs)
    return response.choices[0].message.content

@st.cache_data(ttl=3600, show_spinner=False)
def buscar_fontes_autoridade(palavra_chave):
    """Busca links de alta autoridade (gov, pesquisas) para Deep Links seguros."""
    if not SERPAPI_KEY:
        return "Sem chave Serper configurada."
    url = "https://google.serper.dev/search"
    # Query focada em dados neutros para evitar blogs de concorrentes
    query = f"{palavra_chave} (dados OR estatística OR MEC OR INEP OR pesquisa OR IBGE)"
    payload = json.dumps({"q": query, "gl": "br", "hl": "pt-br", "num": 5})
    headers = {'X-API-KEY': SERPAPI_KEY, 'Content-Type': 'application/json'}
    try:
        response = requests.request("POST", url, headers=headers, data=payload)
        dados = response.json()
        fontes = "🔗 FONTES DE AUTORIDADE EXTERNAS ENCONTRADAS (USE COMO DEEP LINKS SEGUROS):\n"
        if "organic" in dados:
            for idx, res in enumerate(dados["organic"]):
                titulo = res.get('title', 'Sem Título')
                link = res.get('link', '')
                snippet = res.get('snippet', 'Sem resumo')
                
                # LISTA DE RIVAIS EXTERNOS (Removido saseducacao, geekie e outras da Arco)
                rivals_externos = ['poliedro', 'anglo', 'bernoulli', 'objetivo', 'eleva', 'fariasbrito', 'aridesa', 'fibonacci']
                
                if any(rival in link.lower() for rival in rivals_externos):
                    continue
                    
                fontes += f"- FONTE {idx+1}: {titulo}\n  URL EXATA: {link}\n  CONTEXTO: {snippet}\n\n"
        return fontes
    except Exception as e:
        return f"Erro ao buscar fontes externas: {e}"

@st.cache_data(ttl=3600, show_spinner=False)
def buscar_baseline_llm(palavra_chave):
    system_prompt = "Você é um pesquisador de IA sênior. Forneça a resposta que uma IA daria hoje para o termo pesquisado, citando o consenso atual."
    user_prompt = f"O que você sabe sobre: '{palavra_chave}'? Retorne um resumo profundo de como esse tema é respondido atualmente pelas IAs."
    try:
        return chamar_llm(system_prompt, user_prompt, model="openai/gpt-4o-mini", temperature=0.1)
    except Exception as e:
        return f"Erro ao buscar Baseline de IA: {e}"

@st.cache_data(ttl=3600, show_spinner=False)
def buscar_artigos_relacionados_wp(palavra_chave, wp_url, wp_user, wp_pwd):
    """
    RAG Reverso dinâmico com Fallback: Tenta buscar pela keyword. 
    Se o WP não encontrar, puxa os últimos posts para garantir a linkagem interna.
    """
    if not (wp_url and wp_user and wp_pwd):
        return "Sem credenciais do WordPress configuradas para esta marca. Pule a linkagem interna."
    
    import base64
    wp_pwd_clean = wp_pwd.replace(" ", "").strip()
    credenciais = f"{wp_user}:{wp_pwd_clean}"
    token_auth = base64.b64encode(credenciais.encode('utf-8')).decode('utf-8')
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
        'Accept': 'application/json, text/plain, */*',
        'Authorization': f'Basic {token_auth}',
        'Connection': 'keep-alive',
        'Accept-Encoding': 'gzip, deflate, br'
    }

    separador = "&" if "?" in wp_url else "?"
    
    # TENTATIVA 1: Busca nativa do WP pela palavra-chave
    search_url = f"{wp_url}{separador}search={urllib.parse.quote(palavra_chave)}&per_page=6&_fields=id,title,link,excerpt"
    
    try:
        response = requests.get(search_url, headers=headers, timeout=15)
        posts = []
        
        if response.status_code == 200:
            posts = response.json()
            
        # ==========================================
        # TENTATIVA 2: O FALLBACK (O SEGREDO DO REVISOR)
        # Se a busca literal falhar, puxa os 10 últimos posts
        # ==========================================
        if not posts or len(posts) == 0:
            fallback_url = f"{wp_url}{separador}per_page=10&status=publish&_fields=id,title,link,excerpt"
            fallback_res = requests.get(fallback_url, headers=headers, timeout=15)
            if fallback_res.status_code == 200:
                posts = fallback_res.json()

        if not posts:
            return "Nenhum artigo interno altamente relacionado encontrado."
        
        contexto_interno = "🔗 ARTIGOS DO PRÓPRIO BLOG (RAG REVERSO E REFERÊNCIA DE TOM DE VOZ):\n"
        for p in posts[:10]: # Trava de segurança para não estourar os tokens da IA
            titulo = p.get("title", {}).get("rendered", "Sem título")
            link = p.get("link", "")
            
            trecho = p.get("excerpt", {}).get("rendered", "")
            trecho_limpo = re.sub(r'<[^>]+>', '', trecho).strip().replace('\n', ' ')
            
            contexto_interno += f"- Título: {titulo}\n  URL: {link}\n  Trecho do estilo de escrita: '{trecho_limpo}'\n\n"
        
        return contexto_interno
        
    except Exception as e:
        return f"Falha ao conectar com WP da marca para RAG Reverso: {e}"

@st.cache_data(ttl=3600, show_spinner=False)
def buscar_artigos_relacionados_drupal(palavra_chave, d_url, d_user, d_pwd):
    if not (d_url and d_user and d_pwd): return "Sem credenciais Drupal."
    import base64
    import urllib.parse
    
    token_auth = base64.b64encode(f"{d_user}:{d_pwd.replace(' ', '').strip()}".encode('utf-8')).decode('utf-8')
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36', 
        'Accept': 'application/vnd.api+json', 
        'Authorization': f'Basic {token_auth}'
    }
    
    # Extrai a base do site (Ex: https://www.saseducacao.com.br)
    parsed_url = urllib.parse.urlparse(d_url)
    base_url = f"{parsed_url.scheme}://{parsed_url.netloc}"
    
    # TENTATIVA 1: Busca pelo Título
    filtro = f"?filter[title-filter][condition][path]=title&filter[title-filter][condition][operator]=CONTAINS&filter[title-filter][condition][value]={urllib.parse.quote(palavra_chave)}&page[limit]=6"
    
    try:
        res = requests.get(f"{d_url}{filtro}", headers=headers, timeout=15)
        posts = []
        if res.status_code == 200:
            posts = res.json().get("data", [])
            
        # ==========================================
        # TENTATIVA 2: FALLBACK DRUPAL
        # ==========================================
        if not posts or len(posts) == 0:
            fallback_url = f"{d_url}?sort=-created&page[limit]=10"
            fallback_res = requests.get(fallback_url, headers=headers, timeout=15)
            if fallback_res.status_code == 200:
                posts = fallback_res.json().get("data", [])

        if not posts: return "Nenhum artigo encontrado no Drupal."
        
        ctx = "🔗 ARTIGOS DO PRÓPRIO BLOG (RAG REVERSO DRUPAL):\n"
        for p in posts[:10]:
            attrs = p.get("attributes", {})
            titulo = attrs.get('title', '')
            
            path_data = attrs.get('path') or {}
            alias = path_data.get('alias', '') if isinstance(path_data, dict) else ""
            
            # O PULO DO GATO: Força a URL a ser absoluta
            if alias and not alias.startswith('http'):
                # Garante que não teremos barras duplas na junção
                link = f"{base_url}/{alias.lstrip('/')}"
            else:
                link = alias
                
            ctx += f"- Título: {titulo}\n  URL: {link}\n"
        return ctx
        
    except Exception as e:
        return f"Erro Drupal RAG: {e}"

@st.cache_data(ttl=3600, show_spinner=False)
def buscar_estilo_especialista(nome_especialista, df_especialistas):
    """Puxa até 3 artigos do especialista via Jina AI para a IA clonar o estilo de escrita."""
    if not nome_especialista: return ""
    
    links = df_especialistas[df_especialistas['Especialista'] == nome_especialista]['Link do Artigo'].tolist()
    import random
    links_selecionados = random.sample(links, min(3, len(links))) # Pega 3 aleatórios para não estourar limite
    
    contexto = f"📚 CLONAGEM DE PERSONA E REFERÊNCIAS: {nome_especialista}\n"
    
    for link in links_selecionados:
        try:
            jina_headers = {'User-Agent': 'Mozilla/5.0', 'X-Return-Format': 'markdown', 'Accept': 'text/plain'}
            res = requests.get(f"https://r.jina.ai/{link}", headers=jina_headers, timeout=12)
            if res.status_code == 200:
                contexto += f"\n--- Artigo Anterior Escrito por {nome_especialista} ---\n"
                contexto += res.text[:1500] + "...\n" # Pega os primeiros 1500 chars (o ouro do tom de voz)
        except Exception:
            pass
            
    return contexto

@st.cache_data(ttl=300, show_spinner=False)
def listar_posts_wp(wp_url, wp_user, wp_pwd):
    """
    Busca os últimos posts do WP para a aba de Revisão e Auditoria usando máscara de Chrome.
    """
    if not (wp_url and wp_user and wp_pwd):
        return []
    
    import base64
    wp_pwd_clean = wp_pwd.replace(" ", "").strip()
    credenciais = f"{wp_user}:{wp_pwd_clean}"
    token_auth = base64.b64encode(credenciais.encode('utf-8')).decode('utf-8')
    
    # Adicionada a mesma máscara do Ping para driblar o WAF do COC
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'application/json',
        'Authorization': f'Basic {token_auth}',
        'Connection': 'keep-alive'
    }

    separador = "&" if "?" in wp_url else "?"
    
    # Removido o 'draft' para evitar erro 401/403 caso a senha de app tenha privilégios reduzidos
    search_url = f"{wp_url}{separador}per_page=15&status=publish&_fields=id,title,content,link"
    
    try:
        # Aumentamos o timeout para 25s, pois puxar 15 posts do COC pode demorar mais que o ping
        res = requests.get(search_url, headers=headers, timeout=25)
        if res.status_code == 200:
            return res.json()
        else:
            print(f"Erro ao listar posts WP: {res.status_code} - {res.text}")
    except Exception as e:
        print(f"Timeout ou erro na requisição WP: {e}")
        pass
        
    return []

@st.cache_data(ttl=300, show_spinner=False)
def listar_posts_drupal(d_url, d_user, d_pwd):
    if not (d_url and d_user and d_pwd): return []
    import base64
    token_auth = base64.b64encode(f"{d_user}:{d_pwd.replace(' ', '').strip()}".encode('utf-8')).decode('utf-8')
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36', 
        'Accept': 'application/vnd.api+json', 
        'Authorization': f'Basic {token_auth}'
    }
    try:
        res = requests.get(f"{d_url}?sort=-created&page[limit]=15", headers=headers, timeout=15)
        if res.status_code == 200:
            posts = res.json().get("data", [])
            
            lista_formatada = []
            for p in posts:
                attrs = p.get("attributes", {})
                titulo = attrs.get("title") or "Sem Título"
                
                # Proteção contra body nulo
                body_data = attrs.get("body") or {}
                conteudo = body_data.get("value", "") if isinstance(body_data, dict) else ""
                
                lista_formatada.append({
                    "id": p.get("id"),
                    "title": {"rendered": titulo},
                    "content": {"rendered": conteudo}
                })
            return lista_formatada
    except Exception as e:
        print(f"Erro no parser do Drupal: {e}")
        pass
    return []

@st.cache_data(ttl=3600, show_spinner=False)
def buscar_artigos_relacionados_webflow(palavra_chave, w_url, w_user, w_pwd):
    """RAG Reverso e Linkagem Interna para Webflow"""
    if not (w_url and w_pwd): return "Sem credenciais Webflow."
    
    headers = {
        "accept": "application/json",
        "authorization": f"Bearer {w_pwd.strip()}"
    }
    
    try:
        res = requests.get(w_url, headers=headers, timeout=15)
        if res.status_code == 200:
            items = res.json().get("items", [])
            
            # Filtro local de fallback, já que o Webflow não tem busca literal via API nativa
            relevantes = [i for i in items if palavra_chave.lower() in i.get('fieldData', {}).get('name', '').lower()]
            if not relevantes:
                relevantes = items[:10] # Fallback
            
            if not relevantes: return "Nenhum artigo encontrado no Webflow."
            
            ctx = "🔗 ARTIGOS DO PRÓPRIO BLOG (RAG REVERSO WEBFLOW):\n"
            for p in relevantes[:10]:
                field_data = p.get('fieldData', {})
                titulo = field_data.get('name', 'Sem Título')
                slug = field_data.get('slug', '')
                
                # Monta a URL base do blog do Isaac
                link = f"https://isaac.com.br/conteudos/{slug}"
                ctx += f"- Título: {titulo}\n  URL: {link}\n"
            return ctx
    except Exception as e:
        return f"Erro Webflow RAG: {e}"
    return "Falha na conexão com Webflow."

@st.cache_data(ttl=300, show_spinner=False)
def listar_posts_webflow(w_url, w_user, w_pwd):
    """Busca posts para a aba de Auditoria e Revisor"""
    if not (w_url and w_pwd): return []
    
    headers = {
        "accept": "application/json",
        "authorization": f"Bearer {w_pwd.strip()}"
    }
    try:
        # Removi o timeout para garantir que o erro volte
        res = requests.get(w_url, headers=headers) 
        
        if res.status_code == 200:
            items = res.json().get("items", [])
            lista_formatada = []
            for p in items:
                field_data = p.get('fieldData', {})
                titulo = field_data.get('name', 'Sem Título') 
                slug = field_data.get('slug', '')
                conteudo = field_data.get('texto', '') 
                
                lista_formatada.append({
                    "id": p.get("id"),
                    "title": {"rendered": titulo},
                    "content": {"rendered": conteudo},
                    "link": f"https://isaac.com.br/blog/{slug}"
                })
            return lista_formatada
        else:
            # O "DEDO-DURO": Vai mostrar o erro exato no Streamlit
            st.error(f"Erro na API do Webflow ({res.status_code}): {res.text}")
            return []
            
    except Exception as e:
        st.error(f"Erro crítico no Python ao conectar com Webflow: {e}")
    return []

def publicar_webflow(titulo, conteudo_html, meta_dict, w_url, w_user, w_pwd):
    """Envia o rascunho (Draft) direto para o CMS do Webflow"""
    headers = {
        "accept": "application/json",
        "content-type": "application/json",
        "authorization": f"Bearer {w_pwd.strip()}"
    }
    
    payload = {
        "isArchived": False,
        "isDraft": True, # Vai como rascunho
        "fieldData": {
            "name": titulo,
            "slug": slugify(titulo),
            "texto": conteudo_html, # AQUI ESTAVA O ERRO: Enviando para o campo "texto"
            "chamada": meta_dict.get("meta_description", "") # Mapeei a meta description para o seu campo "Chamada"
        }
    }
    
    try:
        response = requests.post(w_url, json=payload, headers=headers, timeout=30)
        return response
    except Exception as e:
        class ErrorResponse:
            status_code = 500
            text = f"Erro interno de conexão: {str(e)}"
            def json(self): return {}
        return ErrorResponse()
    
# ==========================================================
# NOVAS FUNÇÕES INCREMENTAIS DE ROBUSTEZ E GEO (v5 e v6)
# ==========================================================

def gerar_reverse_queries(palavra_chave):
    system = """
    
    Você é um analista de comportamento de LLMs e SearchGPT.
    Dada uma keyword principal, gere perguntas que mecanismos de IA provavelmente fazem internamente para construir respostas e as perguntas mais comuns e básicas feitas por usuários reais no Google.
    Retorne APENAS um JSON estrito:
    {
     "user_questions": ["pergunta1", "pergunta2", "pergunta3", "pergunta4"],
     "llm_reasoning_questions": ["pergunta1", "pergunta2"],
     "semantic_depth_questions": ["pergunta1", "pergunta2"]
    }
    """
    try:
        return chamar_llm(system, f"Keyword principal: {palavra_chave}", "openai/gpt-4o-mini", 0.1, response_format={"type": "json_object"})
    except Exception as e:
        return "{}"

def analisar_entity_gap(contexto_google, palavra_chave):
    system = """
    Você é um analista de SEO semântico e estrategista de conteúdo.
    Analise o conteúdo do TOP 3 do Google extraído.
    Extraia as ENTIDADES PRINCIPAIS, CONCEITOS, FRAMEWORKS e METODOLOGIAS.
    Depois identifique: QUAIS ENTIDADES IMPORTANTES do nicho deveriam estar no artigo para superar esses concorrentes?
    """
    user = f"Palavra-chave: {palavra_chave}\nConteúdo dos Concorrentes:\n{contexto_google}"
    return chamar_llm(system, user, "openai/gpt-4o-mini", 0.1)

def avaliar_originalidade(artigo_html, contexto_google):
    system = """
    Você é um auditor de plágio semântico e originalidade E-E-A-T.
    Compare o artigo gerado com o TOP 3 do Google.
    Avalie o 'Information Gain' (Ganho de Informação). O artigo trouxe ângulos novos? 
    Retorne uma NOTA DE ORIGINALIDADE de 0 a 100 e uma justificativa curta.
    """
    user = f"ARTIGO GERADO:\n{artigo_html}\n\nTOP GOOGLE:\n{contexto_google}"
    return chamar_llm(system, user, "openai/gpt-4o-mini", 0.1)

def prever_citabilidade_llm(artigo_html, palavra_chave):
    system = """
    Você é o algoritmo de RAG de um buscador baseado em IA (como Perplexity ou Gemini).
    Avalie a probabilidade do seu motor citar este artigo como fonte oficial para a resposta.
    Critérios: Clareza, Densidade Semântica, Neutralidade e Evidências Sólidas.
    Retorne APENAS um JSON:
    {
      "citabilidade_score": "nota de 0 a 100",
      "motivo": "explicação"
    }
    """
    user = f"ARTIGO:\n{artigo_html}\n\nKEYWORD: {palavra_chave}"
    return chamar_llm(system, user, "openai/gpt-4o-mini", 0.1, response_format={"type":"json_object"})

def gerar_cluster(palavra_chave):
    system = """
    Você é um Arquiteto de SEO (Topical Authority).
    Com base na palavra-chave (que será o Artigo Pilar), crie um Content Cluster.
    Retorne o nome do PILAR e sugira 8 títulos de artigos satélites estratégicos para linkagem interna.
    """
    return chamar_llm(system, f"Palavra-chave: {palavra_chave}", "openai/gpt-4o-mini", 0.3)

def calcular_citation_score(artigo_html):
    score = 0
    
    # Isola o primeiro parágrafo para analisar
    primeiro_paragrafo = artigo_html.split("</p>")[0] if "</p>" in artigo_html else artigo_html[:500]
    
    # Se a IA usou negrito no início (destacando o termo/resposta), ganha os 2 pontos da estrutura Answer-First
    if "<strong>" in primeiro_paragrafo: 
        score += 2 
        
    if "Resumo Estratégico" in artigo_html or "Resumo estratégico" in artigo_html: 
        score += 1
    if "Segundo especialistas" in artigo_html or "Especialistas" in artigo_html: 
        score += 1
    if "Perguntas Frequentes" in artigo_html: 
        score += 1
        
    return f"{score}/5"

def calcular_entity_coverage(artigo_html, entity_gap_text):
    system = """
    Você é um analisador de SEO semântico.
    Compare:
    1) ENTIDADES importantes sugeridas (Entity Gap)
    2) ENTIDADES presentes no artigo
    Retorne um JSON:
    {
      "entity_coverage_score": "0-100",
      "entities_present": [],
      "entities_missing": []
    }
    """
    user = f"ENTIDADES RECOMENDADAS:\n{entity_gap_text}\n\nARTIGO:\n{artigo_html}"
    return chamar_llm(system, user, "openai/gpt-4o-mini", 0.1, response_format={"type":"json_object"})

def simular_llm_retrieval(keyword, artigo_html):
    system = """
    Você simula o processo de recuperação de fontes usado por motores de busca baseados em LLM.
    Dada uma pergunta do usuário e um artigo, avalie se o conteúdo seria selecionado como fonte.
    Considere: clareza, estrutura citável, entidades confiáveis, completude, neutralidade.
    Retorne JSON:
    {
      "retrieval_score": "0-100",
      "chance_de_ser_usado_como_fonte": "baixa | média | alta",
      "motivo": "explicação curta"
    }
    """
    user = f"PERGUNTA DO USUÁRIO:\n{keyword}\n\nARTIGO:\n{artigo_html}"
    return chamar_llm(system, user, "openai/gpt-4o-mini", 0.1, response_format={"type":"json_object"})

def detectar_citation_hijacking(artigo_html):
    system = """
    Analise o artigo e identifique vulnerabilidade a AI Citation Hijacking.
    Citation Hijacking acontece quando outro conteúdo concorrente pode responder melhor ou mais direto à mesma pergunta.
    Avalie: ausência de resposta direta, falta de definição clara, falta de estrutura citável, excesso de narrativa.
    Retorne JSON:
    {
      "risco_hijacking": "baixo | médio | alto",
      "pontos_fracos": [],
      "melhorias_recomendadas": []
    }
    """
    user = f"ARTIGO:\n{artigo_html}"
    return chamar_llm(system, user, "openai/gpt-4o-mini", 0.1, response_format={"type":"json_object"})

def simular_resposta_ai(keyword, artigo_html):
    system = """
    Simule como um motor de busca baseado em IA (SGE/Perplexity) responderia a uma pergunta do usuário usando o artigo fornecido APENAS como fonte.
    Produza a resposta final que o usuário veria.
    Depois avalie: clareza, completude, necessidade de outras fontes.
    Retorne JSON:
    {
      "resposta_simulada": "...",
      "qualidade_resposta": "0-100",
      "precisaria_de_outras_fontes": true | false
    }
    """
    user = f"PERGUNTA:\n{keyword}\n\nFONTE:\n{artigo_html}"
    return chamar_llm(system, user, "openai/gpt-4o-mini", 0.2, response_format={"type":"json_object"})

def executar_revisao_geo_wp(palavra_chave, publico, marca, html_atual):
    df = st.session_state['brandbook_df']
    marca_info = df[df['Marca'] == marca].iloc[0].to_dict()
    url_marca = marca_info.get('URL', '')

    system = """Você é um Revisor Sênior de SEO e Engenheiro de Prompt GEO.
    Sua missão é avaliar um artigo HTML antigo ou mal formatado e reescrevê-lo para atingir a nota máxima nos critérios E-E-A-T e nas heurísticas do Motor GEO.
    
    DIRETRIZES DE REVISÃO E REESCRITA OBRIGATÓRIAS:
    1. ASSIMETRIA VISUAL EXTREMA: Destrua blocos de texto maciços. Intercale parágrafos "maiores" (3-4 linhas) com parágrafos de UMA ÚNICA FRASE (respiro visual profundo). É proibido que os parágrafos tenham tamanho simétrico.
    2. ANSWER-FIRST: Crie um <h2>Resposta rápida para: [palavra-chave]</h2> logo no início e entregue a resposta mastigada em 2 linhas com a tag <p><strong>Resposta direta:</strong>.
    3. CHUNK CITABILITY: Insira um <p><strong>Definição:</strong> com menos de 30 palavras no início. Limite listas (<ul>) a no máximo 2 em todo o artigo.
    4. BRANDBOOK DA MARCA: Reescreva trechos fora de tom usando o Tom de Voz e Posicionamento exigidos no briefing. Garanta que o nome da marca seja linkado para a URL oficial.
    5. PRESERVAÇÃO DE DADOS: Mantenha as informações e ideias do texto original. Não invente "Estudos da OCDE" ou dados matemáticos se eles não estiverem no texto original.
    6. Mantenha os marcadores `<br>Resumo Estratégico<br>` e `<br>Perguntas Frequentes<br>` onde achar pertinente para o novo esqueleto.
    7. PRESERVAÇÃO DE LINKS E IMAGENS (REGRA INTOCÁVEL): É ESTRITAMENTE PROIBIDO remover, alterar URLs, ou deletar tags `<a>` (hiperlinks), `<img>` e `<figure>` que já estão no HTML original. Você deve reposicioná-las logicamente no novo texto, mantendo os atributos `href`, `src` e classes intactos. O seu trabalho é melhorar o copywriting e a estrutura em volta da mídia, NUNCA apagar o trabalho de linkagem interna/externa e imagens que o redator original já fez.
    8. CORREÇÃO DE CAPITALIZAÇÃO (CRÍTICO): Revise todos os títulos (H1, H2, H3). Se eles estiverem em "Title Case" (Todas As Iniciais Maiúsculas), reescreva-os IMEDIATAMENTE para o padrão brasileiro "Sentence Case" (Apenas a primeira letra e nomes próprios em maiúscula).
    
    RETORNE EXCLUSIVAMENTE UM JSON SEGUINDO ESTE FORMATO EXATO:
    {
        "diagnostico": "Resumo curto das falhas originais de SEO/GEO encontradas.",
        "melhorias_aplicadas": ["Melhoria 1", "Melhoria 2"],
        "html_novo": "O código HTML completo reescrito e otimizado"
    }
    """
    
    user = f"""
    PALAVRA-CHAVE FOCO: '{palavra_chave}'
    PÚBLICO-ALVO: {publico}
    MARCA ALVO: {marca}
    URL DA MARCA OBRIGATÓRIA: {url_marca}
    
    DIRETRIZES DA MARCA ({marca}):
    - Posicionamento: {marca_info['Posicionamento']}
    - Tom de Voz Exigido: {marca_info['TomDeVoz']}
    - Regras Positivas: {marca_info.get('RegrasPositivas', '')}
    - Proibido (Regras Negativas): {marca_info['RegrasNegativas']}
    
    TEXTO ORIGINAL PARA AUDITORIA E REESCRITA (HTML):
    {html_atual}
    """
    
    return chamar_llm(system, user, model="anthropic/claude-4.5-sonnet", temperature=0.3, response_format={"type": "json_object"})

# ==========================================================
# NOVAS MÉTRICAS MATEMÁTICAS RAG / GEO (V7.0)
# ==========================================================
def extrair_numero(valor):
    try:
        if isinstance(valor, dict):
            valor = json.dumps(valor)
        match = re.search(r'\d+', str(valor))
        if match:
            return int(match.group())
    except:
        pass
    return 0

def calcular_geo_score_matematico(citation_score, originalidade, citabilidade, entity_coverage_str):
    # Converte tudo para número
    citation = extrair_numero(citation_score) * 20  # Multiplica por 20 para virar escala 0-100
    original = extrair_numero(originalidade)
    cita_llm = extrair_numero(citabilidade)
    
    # Extrai o score de entidades do dict que a IA gerou antes
    try:
        entity_dict = json.loads(entity_coverage_str)
        entity = int(entity_dict.get("entity_coverage_score", 0))
    except:
        entity = extrair_numero(entity_coverage_str)

    # Cálculo Ponderado Matemático (Soma 100%)
    geo = (0.35 * citation) + (0.25 * cita_llm) + (0.25 * entity) + (0.15 * original)

    return {
        "citation_score_normalizado": f"{citation}/100",
        "citabilidade_llm": cita_llm,
        "originalidade": original,
        "entity_coverage": entity,
        "geo_score_final": round(geo, 2),
        "veredito": "Score calculado matematicamente via heurística RAG com pesos fixos (não subjetivo)."
    }

def avaliar_chunk_citability(artigo_html):
    paragrafos = artigo_html.split("</p>")
    definicoes = 0
    listas = artigo_html.count("<li>")
    paragrafos_curtos = 0

    for p in paragrafos:
        texto_limpo = re.sub(r'<[^>]+>', '', p).strip()
        palavras = len(texto_limpo.split())
        if ":" in texto_limpo and palavras < 40 and palavras > 5:
            definicoes += 1
        if 10 < palavras < 35:
            paragrafos_curtos += 1

    # NOVA LÓGICA DE FREIO: Máximo de 15 pontos para listas (aprox. 5 itens no total)
    pontos_lista = min(listas * 3, 15)

    score = (definicoes * 10) + pontos_lista + (paragrafos_curtos * 2)
    score = min(score, 100)
    return {
        "chunk_citability_score": score,
        "definicoes_estrategicas_detectadas": definicoes,
        "itens_de_lista": listas,
        "paragrafos_de_leitura_rapida": paragrafos_curtos
    }

def avaliar_answer_first(artigo_html):
    # Pega apenas o primeiro parágrafo para análise
    match = re.search(r'<p>(.*?)</p>', artigo_html, re.DOTALL)
    if not match:
        return {"answer_first_score": 0, "status": "Erro: Parágrafo inicial não encontrado."}
    
    primeiro_paragrafo = match.group(1).lower()
    
    # Novos padrões semânticos (o que uma resposta direta parece naturalmente)
    padroes_suaves = ["é ", "são ", "refere-se a", "consiste em", "representa", "trata-se de"]
    
    encontrou_padrao = any(p in primeiro_paragrafo for p in padroes_suaves)
    
    # Se a resposta está no primeiro parágrafo (curto) e usa um verbo de definição
    if encontrou_padrao and len(primeiro_paragrafo.split()) < 50:
        return {
            "answer_first_score": 100, 
            "status": "Excelente: Resposta integrada ao primeiro parágrafo.",
            "metodo": "Detecção Semântica"
        }
    
    return {
        "answer_first_score": 40, 
        "status": "Alerta: O primeiro parágrafo parece longo ou não define o termo rapidamente."
    }

def simular_rag_chunks(artigo_html, keyword):
    chunks = artigo_html.split("\n\n")
    resultados = []
    for c in chunks:
        texto_limpo = re.sub(r'<[^>]+>', '', c).strip()
        if not texto_limpo: continue
        score = 0
        palavras = texto_limpo.lower()
        if keyword.lower() in palavras:
            score += 30
        score += palavras.count(keyword.lower()) * 5
        if ":" in texto_limpo: score += 10
        if len(texto_limpo.split()) < 45: score += 10
        resultados.append({"chunk": texto_limpo[:150] + "...", "score": score})
    
    top_chunks = sorted(resultados, key=lambda x: x["score"], reverse=True)[:3]
    return {"top_chunks_para_llm": top_chunks, "retrieval_strength": round(sum([c["score"] for c in top_chunks])/3, 2) if top_chunks else 0}

def calcular_evidence_density(artigo_html):
    texto_limpo = re.sub(r'<[^>]+>', '', artigo_html).strip()
    numeros = len(re.findall(r'\b\d+\b', texto_limpo))
    porcentagens = len(re.findall(r'\d+%', texto_limpo))
    links = artigo_html.count("href=")
    score = min((numeros * 2) + (porcentagens * 5) + (links * 10), 100)
    return {"evidence_density_score": score, "numeros_absolutos": numeros, "porcentagens": porcentagens, "links_de_referencia": links}

def calcular_information_gain(artigo_html, google_ctx):
    palavras_artigo = set(re.findall(r'\w+', re.sub(r'<[^>]+>', '', artigo_html).lower()))
    palavras_serp = set(re.findall(r'\w+', google_ctx.lower()))
    novas = palavras_artigo - palavras_serp
    score = min(len(novas) / 8, 100) # Matemático bruto
    return {"information_gain_score": round(score, 2), "palavras_unicas_trazidas": len(novas)}

def refinar_artigo_html(html_atual, instrucoes):
    """Permite que a IA edite apenas partes específicas de um artigo já gerado."""
    system = """Você é um Revisor Sênior e Editor de HTML.
    Sua tarefa é modificar um artigo HTML existente ESTRITAMENTE de acordo com as instruções do usuário.
    
    REGRAS CRÍTICAS:
    1. APLIQUE APENAS A MUDANÇA SOLICITADA. Não reescreva o tom de voz e não altere partes do texto que não foram mencionadas na instrução.
    2. MANTENHA TODO O CÓDIGO HTML INTACTO. Preserve todas as tags (<h1>, <h2>, <p>, <ul>), links (<a href...>) e imagens (<img>) exatamente como estão, a menos que a instrução peça para alterá-las.
    3. Retorne EXCLUSIVAMENTE o código HTML finalizado e completo. Pare de gerar texto imediatamente após a última tag HTML. Nada de introduções, comentários ou marcações (```html).
    """
    user = f"INSTRUÇÃO DE ALTERAÇÃO:\n{instrucoes}\n\nARTIGO ORIGINAL (HTML):\n{html_atual}"
    
    return chamar_llm(system, user, model="anthropic/claude-4.5-sonnet", temperature=0.2)

import os
import unicodedata

def slugify(text):
    """Transforma 'SAS Educação' em 'sas_educacao' para achar a pasta correta."""
    text = unicodedata.normalize('NFKD', text).encode('ASCII', 'ignore').decode('utf-8')
    return text.lower().replace(' ', '_').replace('-', '_')

@st.cache_data(ttl=3600, show_spinner=False)
def ler_referencias_locais(marca_nome):
    slug = slugify(marca_nome)
    caminho_pasta = os.path.join("referencias_tom", slug)
    
    if not os.path.exists(caminho_pasta):
        return ""
        
    texto_extraido = ""
    try:
        for nome_arquivo in os.listdir(caminho_pasta):
            caminho_arquivo = os.path.join(caminho_pasta, nome_arquivo)
            if os.path.isfile(caminho_arquivo):
                if nome_arquivo.lower().endswith('.txt'):
                    with open(caminho_arquivo, 'r', encoding='utf-8') as f:
                        texto_extraido += f"\n--- {nome_arquivo} ---\n{f.read()[:3000]}\n"
                elif nome_arquivo.lower().endswith('.pdf'):
                    import PyPDF2
                    with open(caminho_arquivo, 'rb') as f:
                        leitor = PyPDF2.PdfReader(f)
                        texto_pdf = ""
                        for pagina in leitor.pages[:5]: # Lê até 5 páginas para não estourar tokens
                            texto_pdf += pagina.extract_text() + "\n"
                        texto_extraido += f"\n--- {nome_arquivo} ---\n{texto_pdf[:3000]}\n"
                elif nome_arquivo.lower().endswith('.docx'):
                    import docx
                    doc = docx.Document(caminho_arquivo)
                    texto_docx = "\n".join([p.text for p in doc.paragraphs])
                    texto_extraido += f"\n--- {nome_arquivo} ---\n{texto_docx[:3000]}\n"
    except Exception as e:
        print(f"Erro ao ler referencias de {marca_nome}: {e}")
        
    return texto_extraido

@st.cache_data(ttl=3600, show_spinner=False)
def sintetizar_voz_gemini(brandbook_texto, conteudos_referencia):
    """Agente 1: Analisa os documentos e cria um Blueprint de Tom de Voz."""
    if not conteudos_referencia or len(conteudos_referencia.strip()) < 50:
        return f"Siga o Brandbook original:\n{brandbook_texto}"

    system = """
    Você é um Analista Chefe de Copywriting. Sua missão é ler um Brandbook oficial e cruzar com exemplos reais de textos já publicados pela marca.
    Extraia um "Manual de Clonagem de Voz" cirúrgico contendo:
    1. Comprimento médio das frases (curtas/rápidas ou longas/acadêmicas).
    2. Nível de formalidade e jargões favoritos usados nos textos.
    3. Como a marca faz transições.
    4. O que a marca NUNCA faz (baseado na ausência de padrões).
    
    Devolva um guia estrito de instruções (máximo 400 palavras) para guiar o Redator (Claude).
    """
    
    user = f"DIRETRIZES DO BRANDBOOK:\n{brandbook_texto}\n\nTEXTOS DE REFERÊNCIA (Aprenda com eles):\n{conteudos_referencia}"
    
    try:
        # Usa o Gemini 2.5 Pro via OpenRouter
        return chamar_llm(system, user, model="google/gemini-2.5-pro", temperature=0.2)
    except Exception as e:
        return f"Erro no Agente Gemini: {e} - Siga o brandbook: {brandbook_texto}"
        
# ==========================================
# 4. MOTOR PRINCIPAL (COM AS TRAVAS E INCREMENTOS)
# ==========================================
def executar_geracao_completa(palavra_chave, marca_alvo, publico_alvo, conteudo_adicional="", conteudo_proprietario="", modo_humanizado=False, especialista_nome=None, instrucao_livre=""):
    df = st.session_state['brandbook_df']
    marca_info = df[df['Marca'] == marca_alvo].iloc[0].to_dict()
    url_marca = marca_info.get('URL', '')
    from datetime import datetime
    ano_atual = datetime.now().year

    # ROTEADOR DE CMS AQUI
    cms_url, cms_user, cms_pwd, cms_type = obter_credenciais_cms(marca_alvo)

    st.write(f"🕵️‍♂️ Fase 0: Buscando Google, IAs e ativando Agente Gemini...")
    
    # Prepara os dados pro Gemini
    brandbook_txt = f"Tom de Voz: {marca_info['TomDeVoz']} | Regras: {marca_info.get('RegrasPositivas', '')} | Proibido: {marca_info['RegrasNegativas']}"
    referencias_locais = ler_referencias_locais(marca_alvo)
    
    # Busca os dados do especialista (Ghostwriting)
    contexto_ghostwriting = ""
    if especialista_nome:
        contexto_ghostwriting = buscar_estilo_especialista(especialista_nome, st.session_state['especialistas_df'])
        
    with concurrent.futures.ThreadPoolExecutor() as executor:
        futuro_google = executor.submit(buscar_contexto_google, palavra_chave)
        futuro_fontes = executor.submit(buscar_fontes_autoridade, palavra_chave) 
        futuro_ia = executor.submit(buscar_baseline_llm, palavra_chave)
        futuro_reverse = executor.submit(gerar_reverse_queries, palavra_chave)
        
        # AGENTE GEMINI ENTRA AQUI
        futuro_gemini = executor.submit(sintetizar_voz_gemini, brandbook_txt, referencias_locais)
        
        # ... rotas wp/drupal mantidas iguais ...
        
        # O script decide qual CMS atacar
        if cms_type == "drupal":
            futuro_wp_rag = executor.submit(buscar_artigos_relacionados_drupal, palavra_chave, cms_url, cms_user, cms_pwd)
        elif cms_type == "webflow":
            futuro_wp_rag = executor.submit(buscar_artigos_relacionados_webflow, palavra_chave, cms_url, cms_user, cms_pwd)
        else:
            futuro_wp_rag = executor.submit(buscar_artigos_relacionados_wp, palavra_chave, cms_url, cms_user, cms_pwd)
        
        try:
            contexto_google = futuro_google.result(timeout=60) # Aumentado
        except concurrent.futures.TimeoutError:
            contexto_google = "Aviso: A busca orgânica demorou muito. Conteúdo ignorado para manter a velocidade."
        try:
            fontes_externas = futuro_fontes.result(timeout=45) 
        except:
            fontes_externas = "Aviso: Sem fontes externas adicionais."
        try:
            baseline_ia = futuro_ia.result(timeout=60) # Aumentado
        except concurrent.futures.TimeoutError:
            baseline_ia = "Aviso: O motor de Baseline demorou muito a responder. Ignorado."
        try:
            reverse_queries = futuro_reverse.result(timeout=20)
        except:
            reverse_queries = "{}"
        try:
            manual_voz_gemini = futuro_gemini.result(timeout=40)
        except concurrent.futures.TimeoutError:
            manual_voz_gemini = f"Timeout Gemini. Use Brandbook: {brandbook_txt}"    
        try:
            contexto_wp = futuro_wp_rag.result(timeout=25) # Aumentado para dar tempo do Firewall do WP responder
        except:
            contexto_wp = "Erro de timeout ao buscar links internos."
        

    st.write("🔍 Fase 0.5: Analisando Entity Gap e Oportunidades Semânticas...")
    entity_gap = analisar_entity_gap(contexto_google, palavra_chave)

    st.write("🧠 Fase 1: Planejamento Editorial (GPT-4o)...")

    system_1 = """
Você é um Estrategista de Conteúdo GEO (LLM + Search) e Editor-Chefe orientado por E‑E‑A‑T.
Objetivo: produzir um briefing que entregue GANHO DE INFORMAÇÃO e fuja de estruturas genéricas.

REGRAS-MESTRAS (obrigatórias):
1) Nada de “definições básicas” ou “o que é”. O leitor já domina fundamentos. Busque ângulos originais e comparativos.
2) Zero jargão vazio. Frases curtas, voz ativa, tom assertivo.
3) Anti-alucinação total: só liste dados/estudos se houver URL pública verificável.
4) Neutralidade competitiva: ignore marcas privadas concorrentes presentes no contexto bruto.
5) Saída sempre em pt-BR.
6)GATILHOS DE VETO E ANTI-ALUCINAÇÃO (TOLERÂNCIA ZERO):
- REGRA DO DADO ÓRFÃO: É TERMINANTEMENTE PROIBIDO criar briefings sugerindo estatísticas exatas (ex: "37% de aumento", "9 em cada 10") a menos que você tenha a URL profunda e exata fornecida no contexto orgânico. Se não tiver a URL de pesquisa empírica, force o redator a focar em "Argumentação Lógica e Qualitativa" e proíba o uso de números absolutos ou percentuais.
- BLINDAGEM E LINK DE MARCA: Oriente o redator a usar a Marca Alvo exatamente como fornecida e a criar um link (href) para a URL Oficial da marca toda vez que ela for mencionada no texto.

ENTREGÁVEIS DO BRIEFING:
A) ÂNGULO NARRATIVO ÚNICO: escolha 1 (ex.: Quebra de Mito; Guia Tático; Análise de Tendência; Framework Operacional). Justifique em 2-3 linhas focado NAS DORES do público-alvo informado.
B) ESTRUTURA ANTI-FÓRMULA (H2): proponha 4 H2 provocativos, específicos e complementares (sem “O que é”, “Benefícios”, “Conclusão”).
C) MAPA DE EVIDÊNCIAS E DEEP LINKS: Vasculhe o contexto orgânico e resgate o MÁXIMO possível de DEEP LINKS. REGRA DE OURO (ANTI-ALUCINAÇÃO): É ESTRITAMENTE PROIBIDO inventar, deduzir ou construir URLs falsas. Você SÓ PODE extrair e recomendar links que já existem LITERALMENTE no texto do contexto bruto. Se não houver URL lá, não recomende nenhuma.
REGRA CRÍTICA E VETO DE CONCORRÊNCIA: É ESTRITAMENTE PROIBIDO extrair ou sugerir links de domínios de sistemas de ensino concorrentes (ex: Poliedro, Anglo, Bernoulli, SAS, Objetivo, Farias Brito, Ari de Sá, Eleva) ou de sites de outras escolas particulares. 
FILTRO ANTI-PUBLIEDITORIAL: Analise a URL antes de sugeri-la. Se a URL contiver termos como "especial-publicitario", "patrocinado", "publieditorial" ou "branded-content" (como o link do G1 que você vê no contexto), É PROIBIDO USÁ-LA. Só extraia links de jornais neutros, pesquisas, MEC, INEP ou portais governamentais.
E) ENTITY AUTHORITY GRAPH: Liste pelo menos 6 entidades institucionais relevantes para o tema para reforçar autoridade semântica.
F) GATILHO DE MARCA (SEM ALUCINAÇÃO): descreva como a marca aparecerá no terço final como um “Estudo de Caso Prático”. FOQUE APENAS na solução específica (o que a plataforma faz/metodologia). É EXPRESSAMENTE PROIBIDO inventar números de clientes (ex: "um grupo de 5 escolas"), inventar taxas de conversão ou cenários fictícios de antes/depois.
G) MAPA DE CONCRETUDE (PROVAS E BENCHMARKS): Vasculhe o contexto em busca de dados REAIS, números absolutos, benchmarks e comparações tangíveis. Se o contexto mencionar um "mini-caso" (ex: "escola em SP aumentou retenção em 18%"), extraia isso para o Redator usar como micro-história. Nada de conceitos abstratos.
"""

    user_1 = f"""
Palavra-chave ou Consulta: '{palavra_chave}'

Público-Alvo Foco Deste Artigo: {publico_alvo}
    
CONTEÚDO ADICIONAL DO ESPECIALISTA (DIRECIONAMENTO HUMANO):
{conteudo_adicional if conteudo_adicional else "Nenhum conteúdo extra fornecido."}

Contexto extraído do Google (Serper + Jina):
Contexto Google: {contexto_google}
DEEP LINKS DE AUTORIDADE: {fontes_externas}

Baseline de IAs (consenso atual):
{baseline_ia}

Reverse Queries (Perguntas de LLMs para estruturar o texto e FAQ):
{reverse_queries}

Marca Alvo: {marca_alvo}
URL da Marca: {url_marca}
- Posicionamento: {marca_info['Posicionamento']}
- Territórios Estratégicos: {marca_info['Territorios']}

Instruções:
- Construa o briefing completo seguindo as REGRAS-MESTRAS e ENTREGÁVEIS.
- Use as Reverse Queries para entender a intenção de busca profunda da IA.
- Se o contexto carecer de dados confiáveis com URL, declare FOCO CONCEITUAL (sem inventar números).
"""

    analise = chamar_llm(system_1, user_1, model="openai/gpt-4o", temperature=0.3)

    st.write("✍️ Fase 2: Redigindo em HTML Avançado (Claude 4 Sonnet)...")

    if modo_humanizado:
        st.write("✨ Modo Empático ativado: Focando em cadência humana e fluidez...")
        system_2 = """
Você é um Especialista em SEO Semântico (GEO) e um profissional de educação/gestão com vasta experiência prática. 
Sua missão é traduzir o Tom de Voz corporativo em um texto que não pareça um artigo de blog encomendado, mas sim um relato ou análise de quem vive a realidade educacional diariamente.

1. A VOZ DA TRINCHEIRA (EXPERIÊNCIA REAL):
- Escreva como alguém que já participou de reuniões tensas com mantenedores, ouviu reclamações de pais e acompanhou dinâmicas reais de sala de aula. 
- Evite a empatia genérica e enlatada de IA (Ex: NUNCA use "Sabemos que gerir uma escola é um desafio..."). Vá direto para o problema real.
- IMPERFEIÇÃO HUMANA CONTROLADA: A cadência deve parecer orgânica. Não tente fechar todos os parágrafos com uma conclusão perfeita ou "redonda". Use transições secas. Textos humanos reais têm cortes e vão direto ao ponto.

2. O DETECTOR DE ROBÔS E CLICHÊS (BLACKLIST ABSOLUTA):
- VETO A FRASES DE IMPACTO GENÉRICAS: Se usar uma frase curta, ela deve trazer INFORMAÇÃO, não drama. Estão TERMINANTEMENTE PROIBIDAS as frases: "Os números não mentem", "Esta é uma falsa dicotomia", "O segredo está em", "Estamos diante de".
- VETO DE VOCABULÁRIO "IA": Jamais use: "no cenário atual", "cada vez mais", "divisor de águas", "é inegável que", "neste artigo veremos", "em resumo", "por fim", "transcendeu".

3. ANCORAGEM NO MUNDO REAL (CRÍTICO PARA INDETECTABILIDADE):
- MEMÓRIA OPERACIONAL: Sempre que explicar um conceito teórico, obrigatoriamente "encoste a ideia na realidade" com uma micro-cena plausível (Ex: um professor lidando com celulares no fundo da sala, uma decisão financeira no meio do semestre, o momento da renovação de matrículas).
- DETALHE IMPERFEITO: Inclua pequenos detalhes contextuais em suas explicações que não são essenciais para o argumento central, mas gritam "fator humano" (ex: "em turmas mais agitadas", "na primeira semana de provas", "quando o sistema trava").
- FRICÇÃO ANALÍTICA: O texto não pode ser um mar de positividade. Em pelo menos um H2, questione uma prática comum do mercado, aponte um efeito colateral inesperado ou discorde do senso comum. Mostre atrito intelectual.

4) PARÁGRAFO DE IMPACTO (ANSWER-FIRST INTEGRADO): O artigo deve começar obrigatoriamente com um parágrafo de no máximo 4 linhas que combine a introdução com a resposta direta à intenção de busca. 
- É ESTRITAMENTE PROIBIDO usar cabeçalhos como "Resposta rápida para:" ou etiquetas como "Resposta direta:".
- O texto deve ser fluido: comece definindo o conceito e entregando a solução logo nas primeiras duas frases. 
- Use negrito (<strong>) na palavra-chave principal e na parte mais importante da resposta para destacar a densidade semântica para os motores de busca.

5. REGRAS DE LINKAGEM E BLINDAGEM E-E-A-T (TOLERÂNCIA ZERO):
- VETO TOTAL A RIVAIS: É ESTRITAMENTE PROIBIDO citar o nome ou link de QUALQUER outra escola privada ou sistema de ensino concorrente no Brasil (ex: Balão Vermelho, Anglo, Bernoulli). Ignore-os se aparecerem na pesquisa. A única marca privada permitida é a [Marca Alvo].
- LINK DA MARCA: Sempre que citar a [Marca Alvo], transforme-a num link HTML OBRIGATÓRIO: <a href="[URL_DA_MARCA]" target="_blank">[NOME_DA_MARCA]</a>.
- RASTREABILIDADE (DEEP LINKS): Use os links externos fornecidos no briefing (MEC, OCDE, Portais de Notícias). Ancore-os naturalmente. Se não tiver a URL real fornecida no briefing para um dado/pesquisa, NÃO cite a instituição ou os números. Evite alucinação de fontes.
- RAG REVERSO (LINKS INTERNOS): Você receberá "ARTIGOS INTERNOS DISPONÍVEIS". É uma exigência técnica inegociável inserir hiperlinks <a> para 1 ou 2 desses artigos no meio do seu texto, de forma natural.

6. DIRECIONAMENTO E HTML:
- BÚSSOLA DO ARTIGO: Absorva o bloco "Conteúdo Adicional" (teorias, autores). Expanda esses elementos com seu conhecimento interno, aplicando a memória operacional e a fricção analítica descritas acima.
- ESTUDO DE CASO: Ao falar da solução da [Marca Alvo], não faça um texto de vendas. Mostre o contexto operacional de como a ferramenta/método deles destravou um problema.
- REGRAS TÉCNICAS: Use APENAS <h1>, <h2>, <h3>, <p>, <ul>, <ol>, <li>, <strong>, <a>. O <h1> DEVE TER NO MÁXIMO 60 CARACTERES. Títulos em "Sentence case" (Maiúscula só no início).

Finalize o texto com um corte seco ou uma última reflexão técnica. É rigorosamente proibido usar parágrafos de conclusão clichês. Pare de gerar texto imediatamente após fechar a última tag HTML.
"""
    else:
        st.write("⚙️ Modo GEO Restrito ativado: Focando em compliance estrutural...")
        system_2 = """        
Você é Especialista em SEO Semântico (GEO), Copywriter Sênior e Redator de Autoridade E‑E‑A‑T.
Sua missão é traduzir o Tom de Voz corporativo em um texto altamente engajador, focando cirurgicamente nas dores e aspirações do público-alvo.

MANIFESTO ANTI-ROBÔ E ESTILO DA MARCA:
1) Incorpore RIGOROSAMENTE o Tom de Voz e a essência da marca informada.
1.2) Fale DIRETAMENTE com o Público-Alvo definido. Entenda a realidade deles (ex: um gestor busca eficiência; pais buscam segurança).
1.3) Ritmo, profundidade e elegância. Voz ativa. Evite enchimento.
2) PROIBIDO usar jargões de IA como: "No cenário atual", "Cada vez mais", "É inegável que", "É importante ressaltar", "Neste artigo veremos", "Em resumo", "Por fim". 
2.1) VETO DE VOCABULÁRIO IA E CORPORATIVO (BLACKLIST ABSOLUTA): Estão permanentemente banidas expressões robóticas e advérbios longos: "significativamente", "extremamente", "primeiramente", "foi estruturado para oferecer" (use apenas "oferece"), "tanto X quanto para Y" (use apenas "e"). Banido também: "cenário em transformação", "verdadeiro divisor de águas", "influenciar o desempenho agregado das escolas" (o foco é sempre no ALUNO!).
2.2) CONCRETUDE OBRIGATÓRIA (ZERO FRASES VAZIAS): É proibido fazer afirmações genéricas (ex: "a redação é decisiva") sem justificá-las com DADOS ou LÓGICA RÁPIDA (ex: "...visto que instituições atribuem peso 2 à nota"). Fuja de abstrações.
- Exemplo Ruim: "Sistemas de ensino compartilham uma característica: personalização em escala."
- Exemplo Bom: "Escolas que mais aprovam não ensinam todo mundo igual. Elas identificam onde cada aluno trava antes que o problema vire reprovação." 
- Exemplo Bom 2: "Hoje, algumas escolas conseguem prever — com semanas de antecedência — quais alunos têm maior risco de baixo desempenho."
2.3) RITMO ASSIMÉTRICO E HUMANO: O texto NÃO PODE ter parágrafos do mesmo tamanho. Varie o ritmo drasticamente. Misture:
- Parágrafos de 3 a 4 linhas detalhando um processo.
- Frases de impacto isoladas em uma única linha (para dar soco visual).
- Perguntas retóricas diretas ao leitor.
2.4) AUTORIDADE REAL (SEM MULETA GENÉRICA): É ESTRITAMENTE PROIBIDO usar clichês de falsa autoridade como "Segundo especialistas", "Estudos apontam" ou "A neurociência comprova". Isso destrói a credibilidade. Se você não tem o nome do estudo ou do especialista real vindo do briefing, NÃO CITE. Aproprie-se da informação e explique o mecanismo lógico de forma autoral.
2.5) O ESTUDO DE CASO (ANTI-PUBLI): Quando apresentar a solução da [Marca Alvo], não faça panfletagem barata ("A marca X transforma..."). Assuma uma postura madura e editorial. Mostre o "como". Ex: "É por isso que sistemas maduros, como a [Marca Alvo], pararam de entregar apenas cartilhas e passaram a mapear..."
2.6) FAQ HUMANIZADO E DIRETO: O bloco de "Perguntas Frequentes" não pode parecer uma enciclopédia do MEC. Responda como um consultor experiente e pragmático conversando com o gestor. Use inícios de frase francos e diretos como: "Depende. Mas a maioria das escolas...", "Na prática, o primeiro passo é...", "Não necessariamente...".
2.7) PROIBIÇÃO DE MATEMÁTICA FANTASMA (TOLERÂNCIA ZERO): É ESTRITAMENTE PROIBIDO inventar, deduzir ou gerar porcentagens (%), estatísticas precisas ou frações (ex: "aumenta em 32%", "reduz em 40%", "crescimento de 27%"). Se o briefing não lhe forneceu o número exato, escreva de forma puramente QUALITATIVA (ex: "aumenta significativamente a retenção", "reduz o tempo de preparo"). O uso de um único número estatístico inventado reprovará o artigo sumariamente.
3) Não explique o óbvio; entregue leitura avançada.
4) LINK OFICIAL DA MARCA (ANTI-SPAM): A marca alvo e sua URL serão enviadas a você no briefing. Você É OBRIGADO a transformar o nome da marca em um hiperlink (<a href="[URL_AQUI]" target="_blank">) APENAS NA PRIMEIRA VEZ que ela aparecer no texto (geralmente no Estudo de Caso). Nas menções seguintes, escreva o nome da marca como texto puro, sem link, para não configurar spam aos olhos do Google.

GEO (GENERATIVE ENGINE OPTIMIZATION) E CHUNK CITABILITY – REGRAS OBRIGATÓRIAS:
4) INTRODUÇÃO E LINHA FINA: Após o <h1>, crie uma "Linha Fina" (parágrafo curto em <em>) resumindo o artigo. O 1º parágrafo real deve introduzir o problema direto ao ponto, com dados, SEM usar cabeçalhos artificiais como 'Resposta rápida para'.
4.1) FRAMEWORK DE PRODUTO OBRIGATÓRIO (H2): A estrutura do texto deve seguir EXATAMENTE esta ordem narrativa: 
- 1. Introdução e Apresentação do Produto/Dor. 
- 2. Explicação resumida do que é a ferramenta e problema que resolve. 
- 3. Detalhes de Como Funciona (Jornada). 
- 4. Vantagens (sempre após o 'Como funciona'). 
- 5. Exemplos detalhados. 
- 6. H2 de Encerramento (Ex: "Sobre o [Marca]") com CTA direto para o site.
5) RESUMO ESTRATÉGICO: Insira a linha `<br>Resumo Estratégico<br>` após a introdução e crie um <ul> rápido.
7) FRAMEWORK E LEITURA ESCANEÁVEL (CHUNK CITABILITY COM ASSIMETRIA EXTREMA): Transforme seções em frameworks estruturados. O limite MÁXIMO de um parágrafo é de 4 linhas (aprox. 35 palavras). É OBRIGATÓRIO QUEBRAR A SIMETRIA: Intercale parágrafos "maiores" (25 a 35 palavras) com parágrafos de impacto ultracurtos formados por UMA ÚNICA FRASE (8 a 15 palavras). É TERMINANTEMENTE PROIBIDO que os parágrafos tenham o mesmo tamanho visual. LIMITAÇÃO DE LISTAS: Use no máximo 2 a 3 listas (<ul>) em todo o artigo.
8) BLOCO DE AUTORIDADE ORGÂNICO E RASTREÁVEL: É ESTRITAMENTE PROIBIDO usar muletas genéricas de falsa autoridade como "Segundo especialistas", "Estudos apontam" ou "A neurociência diz". 
- SE HOUVER FONTE PROFUNDA (DEEP LINK): Se o briefing forneceu uma URL ESPECÍFICA de uma matéria ou pesquisa, cite o nome da organização de forma natural e OBRIGATORIAMENTE ancore com o link (ex: "Dados recentes do Inep (<a href='.../pesquisa-exata'>) mostram que..."). 
- SE NÃO HOUVER FONTE OU FOR GENÉRICA: Se você NÃO recebeu uma URL, ou se recebeu apenas a homepage raiz (ex: www.mec.gov.br), NÃO cite o nome de nenhuma instituição e não invente um "especialista". Simplesmente afirme o conceito técnico de forma direta, assumindo a autoridade da própria marca que está escrevendo.

REGRAS HTML E FORMATAÇÃO VISUAL (CRÍTICAS E ABSOLUTAS):
9) Use exclusivamente HTML puro: <h1>, <h2>, <h3>, <p>, <ul>, <ol>, <li>, <strong>, <a>. Sem Markdown ou <img>.
10) O primeiro caractere DEVE ser <h1> e o último DEVE ser o fechamento da última tag HTML. O título <h1> DEVE TER NO MÁXIMO 60 CARACTERES. É ESTRITAMENTE PROIBIDO incluir o ano atual (ex: 2025, 2026) no título H1, mantenha-o atemporal.
11) REGRA DE CAPITALIZAÇÃO (SENTENCE CASE): É ESTRITAMENTE PROIBIDO usar "Title Case" nos títulos H1, H2 e H3. Use o padrão gramatical brasileiro: APENAS a primeira letra da frase e nomes próprios/marcas devem ser maiúsculos (Ex: "Como a tecnologia ajuda escolas", NUNCA "Como A Tecnologia Ajuda Escolas").
12) PROIBIDO PARÁGRAFOS SIMÉTRICOS: Verifique o texto antes de entregar. Se você notar que os parágrafos estão visualmente do mesmo tamanho, fragmente-os imediatamente. Obrigatoriamente inclua frases isoladas para criar respiros visuais profundos.
13) VARIAÇÃO HUMANA DE RITMO (OBRIGATÓRIO E EXTREMO):
Humanos não escrevem com ritmo perfeitamente regular. Introduza variação natural drástica:
- Misture frases normais com frases de altíssimo impacto e curtas.
- É OBRIGATÓRIO que a estrutura visual do texto oscile entre blocos maiores e blocos bem curtos.
14) LISTAS COM CONTEXTO E LIMITE: O texto não pode parecer uma apresentação de slides. Se usar uma lista (respeitando o limite máximo de 3 no texto todo), é obrigatório introduzi-la com contexto e concluí-la com forte interpretação analítica.

REGRAS DE LINKAGEM, FONTES E VETOS (E-E-A-T):
15) VETO TOTAL A RIVAIS, OUTRAS ESCOLAS E PUBLIEDITORIAIS (CRÍTICO): É ESTRITAMENTE PROIBIDO citar o nome ou inserir hiperlinks para QUALQUER outra escola privada, colégio ou sistema de ensino concorrente no Brasil. 
- NOMES PROIBIDOS: Poliedro, Anglo, Bernoulli, SAS, Objetivo, Farias Brito, Ari de Sá, Eleva, Fibonacci, etc. 
- VETO DE MÍDIA COMPRADA: Verifique o slug da URL. Se tiver "especial-publicitario", "patrocinado" ou "publi", ignore-o e NÃO USE. 
A única marca privada do setor educacional que pode ser citada é a [Marca Alvo].

16) PROTOCOLO DE RASTREABILIDADE E EXCEÇÃO DE SEGURANÇA (DEEP LINKS): A autoridade depende de referências reais. Extraia links externos (<a href="..." target="_blank">) EXCLUSIVAMENTE do bloco "O QUE A CONCORRÊNCIA DIZ HOJE" ou do "CONTEÚDO ADICIONAL".
- OBRIGAÇÃO CONDICIONAL: Você NÃO TEM LIMITE de links. Pelo contrário, use uma rica mistura de referências. SE o briefing fornecer URLs válidas, você DEVE espalhar de 4 a 8 links externos (ou mais) pelo texto ancorando afirmações, dados, metodologias e leis.
- EXCEÇÃO DE LEIS E METODOLOGIAS: Sempre que citar uma Lei Federal, diretriz do MEC ou uma Metodologia Ativa específica, é mandatório colocar um link de referência (mesmo que seja um portal do governo como planalto.gov.br ou mec.gov.br).

17) DIVERSIDADE DE FONTES E VETO A HOMEPAGES: Valorizamos publicações jornalísticas e acadêmicas de todos os tipos. Contudo, é ESTRITAMENTE PROIBIDO fazer link para homepages genéricas (ex: a página inicial de um jornal ou de um ministério). O link DEVE ser um caminho completo (Deep Link) extraído do briefing para a matéria/pesquisa específica.

18) TOLERÂNCIA ZERO PARA DADOS ÓRFÃOS E URLs ALUCINADAS (CRÍTICO): É TERMINANTEMENTE PROIBIDO citar o nome de QUALQUER instituição, instituto de pesquisa, associação, estudo governamental ou ONG se você não puder ancorar essa citação numa tag <a href> com uma URL REAL e ESPECÍFICA fornecida no briefing.
- TRAVA DE "COPIAR E COLAR": Você só pode usar uma tag <a href="..."> se estiver copiando a URL EXATA fornecida no briefing. É terminantemente proibido inventar, "montar", deduzir ou adivinhar caminhos de URL (ex: criar "/pesquisa-2024/" após um domínio verdadeiro) só para fingir que é um deep link.
- Não existe autoridade sem comprovação. Se não tiver a URL completa fornecida, NÃO CITE O NOME DA INSTITUIÇÃO, do estudo ou da pesquisa. Substitua a menção nominal por percepções empíricas qualitativas universais e não coloque link.
- A regra de alucinação também vale para números: nunca invente estatísticas.
- Exceção: Dados institucionais da própria Marca Alvo não precisam de link.
- ATENÇÃO: NENHUMA instituição tem "passe livre". Se você escrever a frase "<p><strong>Segundo especialistas:</strong> O Ministério da Educação...", você É OBRIGADO a envelopar "Ministério da Educação" com a tag <a href="..."> apontando para a URL real do MEC extraída do briefing. Se o GPT-4o não te deu a URL exata, APAGUE A REFERÊNCIA IMEDIATAMENTE e mude a frase para uma afirmação universal sua.

19) LINKAGEM INTERNA CONTEXTUAL (RAG REVERSO): No final deste prompt, você receberá a lista "ARTIGOS INTERNOS DISPONÍVEIS". 
- REGRA DE OURO: Você DEVE, OBRIGATORIAMENTE, escolher NO MÍNIMO 2 e NO MÁXIMO 4 artigos dessa lista e inserir o link HTML (<a href="[URL_DO_ARTIGO]">) de forma perfeitamente fluida no meio de um parágrafo do seu texto. 
- Adapte a narrativa do texto para que o jargão do link faça sentido. NUNCA use expressões como "Leia também".
- VETO DE ALUCINAÇÃO INTERNA: Se a lista "ARTIGOS INTERNOS DISPONÍVEIS" estiver vazia, disser "Erro" ou "Nenhum artigo", IGNORE ESTA REGRA. Nunca invente uma URL para o blog da marca.

ESTRATÉGIA EDITORIAL, NARRATIVA E VOZ:
22) DIRECIONAMENTO ESTRATÉGICO DO ESPECIALISTA (BÚSSOLA DO ARTIGO): O usuário pode fornecer um bloco de "Conteúdo Adicional" contendo teorias, autores, insumos próprios ou links. Você não precisa fazer um "copia e cola" literal e engessado, mas DEVE usar esses elementos como a base principal da sua argumentação. Use seu conhecimento interno para expandir as teorias ou autores citados, aprofunde os conceitos sugeridos e costure essas referências de forma fluida e inteligente para enriquecer o texto.
23) FRAMEWORK DO ESTUDO DE CASO (P.A.R.): O seu "Estudo de Caso" não pode parecer um panfleto publicitário. Ele deve ser escrito na estrutura Problema (qual dor técnica havia) > Ação da Marca (qual tecnologia exata foi usada) > Resultado (o ganho institucional listado no brandbook). Use o nome comercial da marca.
24) ENTITY SATURATION: Integre naturalmente as entidades mapeadas para provar domínio do nicho.
25) VOZ EDITORIAL DE ANALISTA: Escreva como um analista que observa padrões do setor educacional.
26) OBSERVAÇÃO OPERACIONAL (ANTI-TEXTO GENÉRICO):
-Sempre que explicar um conceito , inclua uma observação concreta da situação ou implementação.
-Evite abstrações vagas. Prefira descrições operacionais.
27) CONTRAPONTO ANALÍTICO (OBRIGATÓRIO EM PELO MENOS 1 H2):
Inclua pelo menos um momento do texto onde uma crença comum do setor é questionada ou refinada.
28) MICRO-ANÁLISE CAUSAL:
Sempre que apresentar um benefício ou prática, explique rapidamente o mecanismo por trás.
29) MICRO-SÍNTESE:
Após alguns blocos analíticos, inclua uma frase curta que consolide a ideia.

[INSTRUÇÃO DE PROCESSAMENTO OBRIGATÓRIA - EXECUÇÃO CRÍTICA]
Para garantir que as regras de linkagem (E-E-A-T) sejam cumpridas sem alucinações, você DEVE, antes de gerar o artigo, criar um bloco <thought_process> (não será lido pelo usuário) contendo o seu planejamento.
Formato obrigatório:
<thought_process>
1. Link interno RAG escolhido: [Escreva aqui a URL exata do bloco ARTIGOS INTERNOS DISPONÍVEIS] (Ancorado na frase: "...")
2. Links externos do briefing que vou usar: [Escreva as URLs exatas, se houver]
3. Confirmo a regra de segurança: Se não há URLs disponíveis no briefing, não citarei nominalmente NENHUMA instituição de pesquisa e não inventarei números.
4. Marca alvo mapeada para linkagem de Estudo de Caso.
</thought_process>

Após fechar a tag </thought_process>, inicie imediatamente o código HTML do artigo com a tag <h1>. 
Lembre-se: Você é OBRIGADO a incluir os marcadores `<br>Resumo Estratégico<br>` e `<br>Perguntas Frequentes<br>`. Abaixo de Perguntas Frequentes, crie 3 perguntas com <h3> e respostas em <p>.
Pare de escrever IMEDIATAMENTE após fechar a última tag HTML. NUNCA gere auto-avaliações ou comentários finais.
"""
# === LÓGICA DO PROMPT LIVRE ===
    bloco_instrucao_livre = ""
    if instrucao_livre and instrucao_livre.strip():
        bloco_instrucao_livre = f"""

INSTRUCAO DIRETA DO USUARIO (PRIORIDADE MAXIMA)
O usuario solicitou o seguinte formato, estrutura e conteudo:
"{instrucao_livre}"

REGRA DE SOBRESCRICAO: Voce DEVE obedecer estritamente aos topicos, perguntas (H2) e ao formato solicitados acima. Esta instrucao substitui qualquer regra de estrutura do briefing anterior. 
No entanto, voce DEVE manter: 
1. A formatacao em HTML puro.
2. O Tom de Voz da marca.
3. A regra de nao alucinar dados sem link.

"""

    user_2 = f"""
{bloco_instrucao_livre}

Palavra-chave ou Consulta: '{palavra_chave}'

DEEP LINKS EXTERNOS (FONTES REAIS):
{fontes_externas}

CONTEXTO TEMPORAL: Ano de {ano_atual}. Não projete o futuro sem evidência. NUNCA insira o ano no título principal (H1) ou no texto a menos que seja um dado histórico.

CONTEÚDO ADICIONAL DO ESPECIALISTA (DIRECIONAMENTO HUMANO OBRIGATÓRIO):
{conteudo_adicional if conteudo_adicional else "Nenhum conteúdo extra fornecido. Siga apenas o briefing."}

CONTEÚDO PROPRIETÁRIO INEGOCIÁVEL (COPIAR E COLAR EXATAMENTE COMO ESTÁ):
{conteudo_proprietario if conteudo_proprietario else "Nenhum conteúdo proprietário exigido."}
ATENÇÃO: Se houver texto no bloco acima, você é OBRIGADO a encontrar um espaço lógico no artigo e transcrever essa frase ou bloco de texto LITERALMENTE, palavra por palavra, sem resumir ou alterar nenhuma vírgula.

O QUE A CONCORRÊNCIA DIZ HOJE:
{contexto_google}

SEU BRIEFING (siga à risca o ângulo e integre o Entity Authority Graph):
{analise}

DIRECIONAMENTO DE COPYWRITING E MARCA:
- Público-Alvo Deste Texto (Foque toda a narrativa neles): {publico_alvo}
- Marca Alvo: {marca_alvo}
- URL da Marca: {url_marca} (OBRIGATÓRIO: Linkar a marca para esta URL sempre que citada).
- Posicionamento: {marca_info['Posicionamento']}
- Territórios: {marca_info['Territorios']}

MANUAL DE CLONAGEM DE VOZ (CRIADO PELO AGENTE GEMINI):
Você é obrigado a escrever o artigo usando exatamente o ritmo, formalidade e regras extraídas abaixo a partir de textos reais da marca:
{manual_voz_gemini}
- Diretrizes OBRIGATÓRIAS: {marca_info.get('RegrasPositivas', '')}
- O que NÃO fazer: {marca_info['RegrasNegativas']}

ARTIGOS INTERNOS DISPONÍVEIS (RAG REVERSO E TOM DE VOZ):
Abaixo estão artigos já publicados no blog da marca. 
OBJETIVO 1 (ESTILO): Leia os trechos para entender e replicar o vocabulário da marca.
OBJETIVO 2 (LINKAGEM CONTEXTUAL): Você É OBRIGADO a escolher 1 ou 2 destes links e inseri-los no meio do seu texto, usando palavras-chave naturais como âncora (tag <a>). NUNCA use "Leia também".
ATENÇÃO ANTI-ALUCINAÇÃO: Se o bloco abaixo disser "Erro", "Timeout" ou "Nenhum artigo", ignore o OBJETIVO 2 e NÃO crie nenhum link interno.
{contexto_wp}

<checklist_de_seguranca_obrigatorio>
1. AVALIAÇÃO ANSWER-FIRST: Você entregou a resposta exata para a dor do leitor logo nas 3 primeiras linhas do texto (usando negrito no conceito principal)? Verifique se você NÃO usou títulos cafonas como "Resposta rápida para:" (isso é proibido).
2. A sua "Definição" tem menos de 30 palavras? (Se tiver mais, reduza agora).
3. ASSIMETRIA VISUAL: Você quebrou os parágrafos corretamente? Há frases isoladas servindo como parágrafos curtos misturadas com parágrafos de 3 linhas? Se o texto estiver um "bloco de tijolo" igual, altere agora.
4. Você usou todas as entidades obrigatórias mapeadas no briefing?
5. VETO A ESCOLAS E RIVAIS: Verifique seu texto e as URLs dos seus links na tag <a href="...">?. Você citou o nome ou o site de ALGUMA OUTRA ESCOLA PRIVADA ou sistema de ensino que não seja a {marca_alvo}? SE SIM, remova imediatamente.
6. O seu "Estudo de Caso" foca na tecnologia/metodologia real da {marca_alvo}? Verifique se você inventou historinha de cliente fictício ou números falsos. Se sim, APAGUE ISSO.
7. CHECK DE DEEP LINKS: Você incluiu pelo menos 2 links externos? Olhe para as URLs dentro do <a href>. Elas são DEEP LINKS reais? Se usou página inicial, substitua IMEDIATAMENTE por um deep link específico ou apague o link.
8. Você garantiu que TODAS as menções à {marca_alvo} contêm o link <a href="{url_marca}">?
8.1 VERIFICAÇÃO DE RAG (CRÍTICO): Escaneie o HTML que você acabou de redigir. Existe alguma tag <a href="..."> apontando para uma das URLs da lista "ARTIGOS INTERNOS DISPONÍVEIS"? Se não houver, VOCÊ DEVE voltar ao texto AGORA, escolher um parágrafo pertinente e inserir o link envelopado em um texto âncora natural.
9. MATEMÁTICA FANTASMA E RASTREABILIDADE (CRÍTICO): Escaneie o seu texto procurando por %, frações ou aumentos numéricos. Responda mentalmente: Você inventou algum dado (ex: "aumenta em 30%", "triplica a retenção") OU usou um número exato sem ancorá-lo imediatamente em um link real <a href="..."> fornecido no briefing? Se a resposta for SIM para qualquer uma das duas, APAGUE O NÚMERO IMEDIATAMENTE. Se o briefing não te deu o número exato junto com a URL de origem, use apenas comparações qualitativas (ex: "escolas reduzem significativamente a evasão").
10. AUDITORIA UNIVERSAL DE FONTES, DEEP LINKS E URLs FALSAS (RISCO DE FALHA CRÍTICA): Procure por QUALQUER menção a instituições (órgãos governamentais, universidades, institutos de pesquisa, ONGs, etc.) no seu texto. Elas estão dentro de uma tag <a href="...">? Agora, faça uma checagem de honestidade: Você COPIOU essa URL exata do briefing ou você INVENTOU um caminho/slug falso só para fingir ser um Deep Link (ex: inventou um "/artigo/" ou "/pdf/")? Se você não tiver o link exato do briefing, se tiver inventado um pedaço da URL, ou se tiver usado apenas uma homepage genérica (ex: gov.br ou usp.br), APAGUE O NOME DA INSTITUIÇÃO E O LINK AGORA MESMO. É terminantemente proibido citar organizações ou leis usando links genéricos, inventados ou sem hiperlink oficial comprobatório.
11. Você analisou o "CONTEÚDO ADICIONAL DO ESPECIALISTA"? O artigo reflete as ideias, autores ou referências sugeridas ali de forma natural e profunda?
12. O seu título <h1> tem menos de 60 caracteres? Conte as letras.
13. CONTEÚDO PROPRIETÁRIO (CRÍTICO): Verifique se foi fornecido algum "CONTEÚDO PROPRIETÁRIO INEGOCIÁVEL". Se sim, procure no seu texto gerado. A frase está EXACTAMENTE igual ao original, sem nenhuma palavra alterada? Se você resumiu ou alterou a frase, corrija agora colando a frase original.
14. DIRETRIZ DE GHOSTWRITING E AUTORIA (CRÍTICO):{contexto_ghostwriting if contexto_ghostwriting else "Nenhum autor específico selecionado. Use o tom da marca padrão."}
ATENÇÃO: Se o bloco acima contiver artigos de um especialista, você assumirá a IDENTIDADE dele. Absorva o vocabulário, o ritmo e o nível de formalidade que ele usa nos artigos fornecidos. Integre o seu conhecimento sobre a palavra-chave com os conceitos que ele costuma defender. 

</checklist_de_seguranca_obrigatorio>

Escreva o ARTIGO FINAL em HTML conforme as regras GEO. 
ATENÇÃO CRÍTICA: Você é OBRIGADO a incluir os exatos marcadores `<br>Resumo Estratégico<br>` no topo e `<br>Perguntas Frequentes<br>` no final do texto. Abaixo de Perguntas Frequentes, crie 3 perguntas e respostas em formato <h3> e <p>.

Pare de escrever IMEDIATAMENTE após fechar a última tag HTML do FAQ. NUNCA gere auto-avaliações, comentários ou textos que comecem com "AI:".
"""
    artigo_html = chamar_llm(system_2, user_2, model="anthropic/claude-4.5-sonnet", temperature=0.45)
    artigo_html = re.sub(r'^```html\n|```$', '', artigo_html, flags=re.MULTILINE).strip()
    artigo_html = re.sub(r'<thought_process>.*?</thought_process>', '', artigo_html, flags=re.DOTALL).strip()

    
    # GUILHOTINA PYTHON: Corta qualquer "auto-avaliação" da IA que venha depois do fechamento do HTML
    if '<' in artigo_html and '>' in artigo_html:
        artigo_html = artigo_html[artigo_html.find('<') : artigo_html.rfind('>') + 1]

    st.write("🛠️ Fase 3: Extraindo JSON e Metadados via Pydantic...")
    schema_gerado = MetadadosArtigo.model_json_schema() if hasattr(MetadadosArtigo, "model_json_schema") else MetadadosArtigo.schema_json()

    system_3 = f"""
Você é especialista em SEO técnico e Schema.org.
Retorne EXCLUSIVAMENTE um JSON puro, válido e COMPATÍVEL com este schema Pydantic:
{json.dumps(schema_gerado, ensure_ascii=False)}

REGRAS CRÍTICAS:
1. NUNCA inclua markdown, comentários ou campos extras.
2. 'title': 45-60 caracteres (otimizado para H1/SEO, sem marca). É ESTRITAMENTE PROIBIDO inserir o ano atual (ex: 2026) neste campo.
3. 'meta_description': 130-150 caracteres (promessa clara + gancho, sem clickbait).
4. 'dicas_imagens': exatamente 2 strings em inglês, MUITO CURTAS E SIMPLES (máximo 1 a 2 palavras, ex.: "classroom", "students"). É ESTRITAMENTE PROIBIDO gerar frases longas.
5. 'schema_faq': JSON-LD FAQPage com @context "[https://schema.org](https://schema.org)", @type "FAQPage" e mainEntity como lista de objetos Question/acceptedAnswer.
    - As perguntas e respostas DEVEM ser extraídas textualmente da seção Perguntas Frequentes presente no HTML fornecido.
    - Se não houver FAQ no HTML, retorne 'schema_faq': {{}}. 

ANTI-CLOAKING E VALIDAÇÃO:
- Proibido inventar perguntas/respostas que não existam no HTML.
- Proibido inventar dados/anos/links no JSON.
- Saída deve conter apenas as chaves: title, meta_description, dicas_imagens, schema_faq.
"""

    user_3 = f"HTML COMPLETO:\n{artigo_html}"

    dicas_json = chamar_llm(system_3, user_3, model="anthropic/claude-4.5-sonnet", temperature=0.1, response_format={"type": "json_object"})

   # MOTOR DUPLO DE IMAGENS (UNSPLASH + FALLBACK POLLINATIONS)
    try:
        json_limpo = dicas_json.strip().removeprefix('```json').removesuffix('```').strip()
        meta_dicas = json.loads(json_limpo)
        termos_busca = meta_dicas.get('dicas_imagens', [])
        UNSPLASH_KEY = st.secrets.get("UNSPLASH_KEY", "")
        
        if isinstance(termos_busca, list):
            for i, termo in enumerate(termos_busca[:2]):
                img_html_pronta = ""
                if UNSPLASH_KEY:
                    # URL LIMPA E DIRETA
                    url = f"https://api.unsplash.com/search/photos?query={urllib.parse.quote(termo)}&client_id={UNSPLASH_KEY}&per_page=1&orientation=landscape"
                    try:
                        res = requests.get(url, timeout=5)
                        if res.status_code == 200:
                            dados_img = res.json()
                            if "results" in dados_img and len(dados_img["results"]) > 0:
                                img_url = dados_img["results"][0]["urls"]["regular"]
                                alt_text = dados_img["results"][0]["alt_description"] or termo
                                img_html_pronta = f'<img src="{img_url}" alt="{alt_text}" style="width:100%; border-radius:8px;" loading="lazy" decoding="async" />'
                    except Exception:
                        pass
                
                if not img_html_pronta:
                    # FALLBACK LIMPO E DIRETA
                    clean_termo = str(termo).replace("'", "").replace('"', '').strip()
                    p_codificado = urllib.parse.quote(clean_termo)
                    base_poll = "https://image.pollinations.ai/prompt/"
                    img_html_pronta = f'<img src="{base_poll}{p_codificado}?width=1024&height=512&nologo=true&model=flux" alt="{clean_termo}" style="width:100%; border-radius:8px;" loading="lazy" decoding="async" />'
                    
                if img_html_pronta:
                    alvo_replace = '<br>Resumo Estratégico<br>' if i == 0 else '<br>Perguntas Frequentes<br>'
                    artigo_html = artigo_html.replace(alvo_replace, f'{img_html_pronta}\n{alvo_replace}', 1)
    except Exception as e:
        st.error(f"Erro ao injetar imagem: {e}") # Mudei para st.error para você ver se falhar

    # CHAMADAS INCREMENTAIS PÓS-REDAÇÃO (GEO PIPELINE COMPLETO)
    st.write("📊 Fase 4: Calculando Originalidade, Citabilidade GEO e Cluster...")
    score_originalidade = avaliar_originalidade(artigo_html, contexto_google)
    citabilidade = prever_citabilidade_llm(artigo_html, palavra_chave)
    cluster = gerar_cluster(palavra_chave)
    citation_score = calcular_citation_score(artigo_html)

    st.write("🧪 Fase 5: Calculando Matrizes RAG e Entity Coverage...")
    entity_coverage = calcular_entity_coverage(artigo_html, entity_gap)
    
    geo_score = calcular_geo_score_matematico(citation_score, score_originalidade, citabilidade, entity_coverage)
    chunk_citability = avaliar_chunk_citability(artigo_html)
    answer_first = avaliar_answer_first(artigo_html)
    rag_chunks = simular_rag_chunks(artigo_html, palavra_chave)
    evidence_density = calcular_evidence_density(artigo_html)
    information_gain = calcular_information_gain(artigo_html, contexto_google)

    st.write("🔬 Fase 6: Simulação de RAG e Citation Hijacking (Motores LLM)...")
    retrieval_simulation = simular_llm_retrieval(palavra_chave, artigo_html)
    hijacking_risk = detectar_citation_hijacking(artigo_html)
    ai_simulation = simular_resposta_ai(palavra_chave, artigo_html)

    return (
        artigo_html, dicas_json, contexto_google, baseline_ia, entity_gap, 
        score_originalidade, citabilidade, cluster, reverse_queries, 
        citation_score, entity_coverage, geo_score, retrieval_simulation, 
        hijacking_risk, ai_simulation, chunk_citability, answer_first, 
        rag_chunks, evidence_density, information_gain, contexto_wp,
        manual_voz_gemini
    )

def publicar_wp(titulo, conteudo_html, meta_dict, wp_url, wp_user, wp_pwd):
    import base64
    
    seo_title = meta_dict.get("title", titulo)
    meta_desc = meta_dict.get("meta_description", "")
    
    # Payload limpo sem scripts
    payload = {
        "title": titulo,
        "content": conteudo_html,
        "status": "draft",
        "meta": {
            "_yoast_wpseo_title": seo_title,
            "_yoast_wpseo_metadesc": meta_desc
        }
    }
    
    wp_pwd_clean = wp_pwd.replace(" ", "").strip()
    credenciais = f"{wp_user}:{wp_pwd_clean}"
    token_auth = base64.b64encode(credenciais.encode('utf-8')).decode('utf-8')
    
   # MÁSCARA ROBUSTA: Disfarce de navegador real (Chrome no Windows)
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'application/json, text/plain, */*',
        'Content-Type': 'application/json',
        'Authorization': f'Basic {token_auth}',
        'Connection': 'keep-alive',
        'Accept-Encoding': 'gzip, deflate, br'
    }
    
    try:
        response = requests.post(wp_url, json=payload, headers=headers, timeout=30)
        return response
    except Exception as e:
        class ErrorResponse:
            status_code = 500
            text = f"Erro interno de conexão: {str(e)}"
            def json(self): return {}
        return ErrorResponse()

def publicar_drupal(titulo, conteudo_html, meta_dict, d_url, d_user, d_pwd):
    import base64
    # Descobre o nome da rota dinamicamente (ex: node--quark_blog)
    node_type = "node--" + d_url.rstrip('/').split('/')[-1] 
    
    payload = {
        "data": {
            "type": node_type,
            "attributes": {
                "title": titulo,
                "body": {"value": conteudo_html, "format": "full_html"},
                "status": False,
                # TENTATIVA DE BYPASS: Preenchendo o campo obrigatório do SAS com um placeholder genérico
                "field_quark_blog_featured_image": "https://images.unsplash.com/photo-1524995997946-a1c2e315a42f" 
            }
        }
    }
    
    token_auth = base64.b64encode(f"{d_user}:{d_pwd.replace(' ', '').strip()}".encode('utf-8')).decode('utf-8')
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36', 
        'Accept': 'application/vnd.api+json', 
        'Content-Type': 'application/vnd.api+json', 
        'Authorization': f'Basic {token_auth}'
    }
    try:
        return requests.post(d_url, json=payload, headers=headers, timeout=30)
    except Exception as e:
        class ErrorRes: 
            status_code = 500
            text = f"Erro interno de conexão: {str(e)}"
            def json(self): return {}
        return ErrorRes()


import PyPDF2
import docx  # Essa linha é obrigatória para ler Word
import io

def extrair_texto_documentos(arquivos_upados):
    """Lê múltiplos arquivos (PDF, DOCX, TXT) e extrai todo o texto."""
    texto_completo = ""
    for arquivo in arquivos_upados:
        nome_arquivo = arquivo.name.lower()
        texto_completo += f"\n\n--- INÍCIO DO DOCUMENTO: {arquivo.name} ---\n"
        
        try:
            if nome_arquivo.endswith('.pdf'):
                leitor = PyPDF2.PdfReader(arquivo)
                for pagina in leitor.pages:
                    texto_extrato = pagina.extract_text()
                    if texto_extrato:
                        texto_completo += texto_extrato + "\n"
            
            elif nome_arquivo.endswith('.docx'):
                doc = docx.Document(arquivo)
                for paragrafo in doc.paragraphs:
                    texto_completo += paragrafo.text + "\n"
                    
            elif nome_arquivo.endswith('.txt'):
                texto_completo += arquivo.read().decode('utf-8') + "\n"
                
        except Exception as e:
            texto_completo += f"\n[Erro ao ler arquivo {arquivo.name}: {e}]\n"
            
        texto_completo += f"\n--- FIM DO DOCUMENTO: {arquivo.name} ---\n"
        
    return texto_completo
    
def executar_adaptacao_documentos(palavra_chave, publico, marca, texto_base_docs, instrucoes_usuario):
    """
    Transforma documentos brutos em um artigo. 
    Se houver instruções, segue-as. Se não houver, age como o 'Gerador de Teaser/Spoiler' para Leads.
    """
    df = st.session_state['brandbook_df']
    marca_info = df[df['Marca'] == marca].iloc[0].to_dict()
    url_marca = marca_info.get('URL', '')

    # ==========================================
    # LÓGICA DE ROTEAMENTO (TEASER VS CUSTOMIZADO)
    # ==========================================
    if instrucoes_usuario and instrucoes_usuario.strip():
        # MODO 1: SÍNTESE CUSTOMIZADA (O usuário deu a regra)
        comportamento_alvo = f"""
        OBJETIVO ESTRATÉGICO DO USUÁRIO:
        {instrucoes_usuario}
        
        REGRA DE ORDENAÇÃO (FRAMEWORK DE PRODUTO INEGOCIÁVEL):
        ATENÇÃO CRÍTICA: Mesmo que a instrução do usuário acima peça tópicos ou H2 em uma ordem específica, você É OBRIGADO a reorganizar a estrutura para seguir ESTRITAMENTE a seguinte ordem narrativa:
        
        1. Parágrafo de introdução e apresentação.
        2. Explicação resumida do produto, o problema que resolve e as VANTAGENS associadas (É estritamente proibido jogar as vantagens para o final do texto. Elas devem subir e ficar logo após a introdução).
        3. Detalhes de "Como funciona" (pode conter vários parágrafos e a jornada do usuário).
        4. Mais exemplos, dados e detalhes de implementação.
        5. Parágrafo final de amarração.
        6. FECHAMENTO OBRIGATÓRIO: Encerre o artigo criando um último <h2> sobre a marca "{marca}" (Ex: "Sobre o {marca}") contendo um CTA chamando o leitor para acessar o site oficial.
        """
    else:
        # MODO 2: TEASER E CAPTAÇÃO DE LEADS (Padrão se o prompt estiver vazio)
        comportamento_alvo = """
        OBJETIVO ESTRATÉGICO: ESTRUTURA TEASER (A TÉCNICA DO SPOILER)
        O objetivo deste artigo NÃO é entregar todo o conteúdo dos documentos, mas sim gerar curiosidade e atuar como uma página de atração para que o leitor baixe o material completo.
        - É expressamente PROIBIDO resumir todos os tópicos ou listar todas as perguntas/respostas do material. 
        - Faça uma introdução sobre o cenário e escolha APENAS UM conceito forte ou UMA pergunta com resposta do material (que faça sentido para o Território da Marca) para dar como "spoiler" gratuito. Apele para a curiosidade sobre o que ficou de fora.
        - O GATILHO PARA O DOWNLOAD (TOM CONVIDATIVO): No final do texto, crie a transição para o download. Use este framework mental para a chamada: "Quer saber mais sobre quais são os outros pilares/pontos de [Tema] e como isso impacta a sua realidade? Baixe o material completo para receber direcionais práticos..."
        - PLACEHOLDER DO TIME DE GROWTH: Logo após o convite para baixar, insira EXATAMENTE esta tag HTML:
          <div style='background-color: #f3f4f6; padding: 20px; text-align: center; border-radius: 8px; margin-top: 20px;'><strong>[Formulário de Captura do Material inserido pelo time de Growth]</strong></div>
        """

    # ==========================================
    # SYSTEM PROMPT BLINDADO (REGRAS DO SYSTEM_2 INCLUÍDAS)
    # ==========================================
    system = f"""Você é um Copywriter Especialista, Arquiteto de Informação e Engenheiro de Conteúdo GEO.
    Sua missão é ler os documentos brutos fornecidos e construir um novo material HTML.

    COMPORTAMENTO DEFINIDO PARA ESTA TAREFA:
    {comportamento_alvo}

    REGRAS INVIOLÁVEIS DE CONSTRUÇÃO E E-E-A-T (FONTE DA VERDADE):
    1. ANTI-ALUCINAÇÃO ABSOLUTA: Use EXCLUSIVAMENTE as informações, dados, leis e exemplos presentes nos documentos fornecidos. Não invente funcionalidades, estatísticas ou conceitos de fora. A autoridade deste texto deriva ÚNICA E EXCLUSIVAMENTE dos documentos anexados.
    2. ZERO LINKAGEM EXTERNA: Como este material é construído a partir de documentação interna, É ESTRITAMENTE PROIBIDO inventar links externos ou citar URLs da web (como MEC, OCDE, portais de notícias).

    MANIFESTO ANTI-ROBÔ E ESTILO DA MARCA:
    3. DIFERENCIAÇÃO EXTREMA DE MARCA: O seu texto DEVE ser guiado 100% pelo Posicionamento e Territórios da Marca Alvo. 
    4. BRAND WEAVING (INSERÇÃO NATURAL DA MARCA): Integre o nome da marca, seus diferenciais e seu propósito no MEIO do texto. A autoridade e a história da marca devem estar costuradas na narrativa. É OBRIGATÓRIO transformar a primeira menção da marca em um link: <a href="{url_marca}" target="_blank">[NOME DA MARCA]</a>.
    5. BLACKLIST DE IA E VETO DE CÓPIA (TOLERÂNCIA ZERO): É ESTRITAMENTE PROIBIDO usar advérbios terminados em "mente" (ex: significativamente, extremamente), jargões corporativos (ex: "influenciar o desempenho agregado da escola", "escola parceira", "da coleção") ou locuções passivas como "foi estruturado para oferecer" (use o verbo direto: "oferece"). O foco de todo o benefício do texto DEVE SER O ALUNO.
    ATENÇÃO CRÍTICA: Mesmo que essas palavras proibidas estejam escritas literalmente no documento base fornecido, VOCÊ É OBRIGADO A REESCREVÊ-LAS E CORTÁ-LAS. A Blacklist tem prioridade absoluta sobre a fidelidade ao texto original. NUNCA use o H2 "Resposta rápida para:".
    6. CONCRETUDE OBRIGATÓRIA E REESCRITA DE CLICHÊS: Se o documento base fizer afirmações vazias (ex: "A redação é decisiva no ENEM"), você NÃO DEVE agir como um papagaio e apenas copiá-la. Você deve reescrevê-la ou enriquecê-la com fatos lógicos universais (ex: "visto que muitas instituições atribuem peso 2 ou 3 à nota final") para que a frase tenha peso. É proibido gerar introduções vazias.

    GEO E CHUNK CITABILITY (HTML E ESTRUTURA VISUAL):
    7. INTRODUÇÃO E LINHA FINA: O texto DEVE começar com o <h1>. Logo abaixo, crie uma "Linha Fina" (parágrafo em <em>) resumindo o texto. O 1º parágrafo normal deve introduzir a dor/solução direto ao ponto.
    7.1. FRAMEWORK DE PRODUTO: Siga estritamente a ordem de H2: Introdução -> O que é a ferramenta/dor resolvida -> Como Funciona (Detalhes) -> Vantagens -> Exemplos -> CTA final sobre a marca.
    8. ASSIMETRIA VISUAL EXTREMA (CRÍTICO): É TERMINANTEMENTE PROIBIDO que os parágrafos tenham o mesmo tamanho. Intercale parágrafos "maiores" (3 a 4 linhas) com frases de impacto isoladas em uma única linha. O ritmo visual deve oscilar drasticamente.
    9. REGRA DE CAPITALIZAÇÃO (SENTENCE CASE): É ESTRITAMENTE PROIBIDO usar "Title Case" nos títulos H1, H2 e H3. Use o padrão brasileiro: APENAS a primeira letra da frase e nomes próprios devem ser maiúsculos. O H1 deve ter no máximo 60 caracteres.
    10. PREVENÇÃO DE ERRO JSON (CRÍTICO): Seu retorno será processado por json.loads(). É OBRIGATÓRIO usar aspas simples (') nas tags HTML (ex: <h2 class='titulo'>) em vez de aspas duplas. Se precisar usar aspas duplas no texto, coloque uma barra invertida antes da aspa.

    RETORNE EXCLUSIVAMENTE UM JSON:
    {{
        "diagnostico": "Explique brevemente a sua estratégia, como usou os documentos, qual spoiler escolheu (se aplicável) e como aplicou a Assimetria Visual.",
        "melhorias_aplicadas": ["Estruturação sob medida", "Assimetria Visual Aplicada", "Brand Weaving", "Veto a jargões de IA respeitado"],
        "html_novo": "O código HTML completo usando aspas simples e escapando aspas duplas internas."
    }}
    """
    
    user = f"""
    PALAVRA-CHAVE FOCO: '{palavra_chave}'
    PÚBLICO-ALVO: {publico}
    MARCA ALVO: {marca}
    
    DIRETRIZES DA MARCA ({marca}):
    - Posicionamento: {marca_info['Posicionamento']}
    - Territórios Estratégicos: {marca_info.get('Territorios', 'Educação')}
    - Tom de Voz Exigido: {marca_info['TomDeVoz']}
    - Regras Positivas: {marca_info.get('RegrasPositivas', '')}
    - Proibido (Regras Negativas): {marca_info['RegrasNegativas']}
    
    BASE DE CONHECIMENTO (DOCUMENTOS FORNECIDOS PARA SÍNTESE):
    {texto_base_docs}
    """
    
    return chamar_llm(system, user, model="anthropic/claude-4.5-sonnet", temperature=0.3, response_format={"type": "json_object"})
    
# ==========================================
# 5. INTERFACE PRINCIPAL
# ==========================================

if st.session_state['current_page'] == "BrandBook":
    st.markdown("### 🏢 Edite as regras, marcas e diretrizes:")
    st.session_state['brandbook_df'] = st.data_editor(st.session_state['brandbook_df'], num_rows="dynamic", width="stretch")
    st.info("💡 Dica: Adicione regras específicas na coluna 'RegrasPositivas'.")
    
    st.markdown("---")
    st.markdown("### ✍️ Especialistas e Autores (Ghostwriting)")
    st.caption("Adicione o nome do especialista e o link de um artigo escrito por ele (LinkedIn, Blog, etc). O Motor lerá os links para clonar o tom de voz do autor.")
    st.session_state['especialistas_df'] = st.data_editor(st.session_state['especialistas_df'], num_rows="dynamic", width="stretch", key="editor_esp")

elif st.session_state['current_page'] == "Gerador de Artigos":
    gerar_btn = False # <--- ISTO AQUI MATA O NAMEERROR PARA SEMPRE!
    
    if not st.session_state['show_inputs']:
        # ==========================================
        # HERO SECTION E CARDS DE VENDA (SÓ MOSTRA ANTES DO CLIQUE)
        # ==========================================
        st.markdown("""
        <style>
        /* Animação CSS pura para alternar os nomes das IAs sem usar Javascript */
        .animated-ia::after {
            content: 'Gemini';
            animation: switchWord 6s infinite;
            color: #418EDE;
            font-weight: 700;
        }
        @keyframes switchWord {
            0%, 20%  { content: 'Gemini'; }
            25%, 45% { content: 'Perplexity'; }
            50%, 70% { content: 'Chat GPT'; }
            75%, 95% { content: 'Copilot'; }
            100%     { content: 'Gemini'; }
        }
        </style>
        
        <div style="text-align: center; margin-top: -3.5rem; margin-bottom: 1rem;">
            <div class="arco-tag" style="margin-bottom: 0.5rem;">MOTOR DE INTELIGÊNCIA</div>
            <h1 style="font-size: 3rem; margin-top: 0rem; margin-bottom: 0.5rem;">Motor GEO v7.0 <span style="color: #F05D23;">AI Search Native</span></h1>
            <p style="font-size: 1.1rem; color: #4B5563; margin-top: 0; font-family: 'Inter', sans-serif;">
                O objetivo dessa ferramenta é criar artigos otimizados para IAs como <span class="animated-ia"></span>
            </p>
        </div>
        """, unsafe_allow_html=True)

        # BOTÃO PRETO CENTRALIZADO
        col_cta1, col_cta2, col_cta3 = st.columns([1, 1, 1])
        with col_cta2:
            if st.button("Gerar artigo 🚀", type="primary", use_container_width=True):
                st.session_state['show_inputs'] = True
                st.rerun()
        
        # PIPELINE EMBAIXO DO BOTÃO
        st.markdown(pipeline_html, unsafe_allow_html=True)

        # === NOVO: SETA ANIMADA INDICANDO SCROLL ===
        st.markdown("""
        <style>
        @keyframes bounce-down {
            0%, 100% { transform: translateY(0); opacity: 0.4; }
            50% { transform: translateY(10px); opacity: 1; }
        }
        .scroll-indicator {
            text-align: center;
            margin-top: 30px;
            margin-bottom: -10px;
            color: #9CA3AF;
            animation: bounce-down 2s infinite ease-in-out;
        }
        </style>
        <div class="scroll-indicator">
            <svg width="32" height="32" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
                <polyline points="6 9 12 15 18 9"></polyline>
            </svg>
            <div style="font-size: 12px; margin-top: 2px; font-family: 'Inter', sans-serif; font-weight: 500; letter-spacing: 0.05em; text-transform: uppercase;">VEJA MAIS</div>
        </div>
        """, unsafe_allow_html=True)

        # CARDS SELLING LLMS (COM ÍCONES NATIVOS À PROVA DE FALHAS)
        st.markdown("<h3 style='margin-top: 3rem; font-size: 1.5rem;'>As novidades. Veja o que acabou de chegar.</h3>", unsafe_allow_html=True)
        c1, c2, c3, c4, c5 = st.columns(5)
        
        with c1:
            st.markdown("""
            <div class="saas-card">
                <img src="https://upload.wikimedia.org/wikipedia/commons/8/8a/Google_Gemini_logo.svg" alt="Gemini Logo" style="height: 32px; margin-bottom: 16px;">
                <div class="card-title">Refino de Marca (Gemini)</div>
                <div class="card-text">Lê seus materiais de referência, o OURO da marca, e treina o motor para escrever com a sua voz exata.</div>
            </div>
            """, unsafe_allow_html=True)
        with c2:
            st.markdown("""
            <div class="saas-card">
                <img src="https://upload.wikimedia.org/wikipedia/commons/0/04/ChatGPT_logo.svg" alt="GPT-4o Logo" style="height: 32px; margin-bottom: 16px;">
                <div class="card-title">Estrategista (GPT-4o)</div>
                <div class="card-text">Analisa a concorrência e cria briefing com regras E-E-A-T.</div>
            </div>
            """, unsafe_allow_html=True)
        with c3:
            st.markdown("""
            <div class="saas-card">
                <img src="https://commons.wikimedia.org/wiki/Special:FilePath/Claude_AI_symbol.svg" alt="Claude 3.7 Logo" style="height: 32px; margin-bottom: 16px;">
                <div class="card-title">Redator (Claude 4)</div>
                <div class="card-text">A inteligência mais avançada para Copywriting.</div>
            </div>
            """, unsafe_allow_html=True)
        with c4:
            st.markdown("""
            <div class="saas-card">
                <div style="font-size: 2rem; margin-bottom: 8px;">🌐</div>
                <div class="card-title">AI Search Native</div>
                <div class="card-text">Extrai o conteúdo da web em tempo real para mapear 'Entity Gap'.</div>
            </div>
            """, unsafe_allow_html=True)
        with c5:
            st.markdown("""
            <div class="saas-card">
                <div style="font-size: 2rem; margin-bottom: 8px;">🔗</div>
                <div class="card-title">RAG Reverso (WP)</div>
                <div class="card-text">Faz linkagem interna com os artigos que você já publicou.</div>
            </div>
            """, unsafe_allow_html=True)

    else:
        # ==========================================
        # FORMULÁRIO DE GERAÇÃO (MOSTRADO APÓS CLICAR NO BOTÃO)
        # ==========================================
        st.markdown("""
        <div style="text-align: center; margin-bottom: 2rem;">
            <h1 style="font-size: 2rem;">Motor GEO v7.0 <span style="color: #F05D23; font-size: 1.2rem;">AI Search Native</span></h1>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(pipeline_html, unsafe_allow_html=True)

        caixa_topo = st.container()
        st.markdown("<br>", unsafe_allow_html=True) 
        
        col1, col2 = st.columns([1, 2])
        with col1:
            marca_selecionada = st.selectbox("Selecione a Marca", st.session_state['brandbook_df']['Marca'].tolist())
            
            try:
                publicos_da_marca = st.session_state['brandbook_df'][st.session_state['brandbook_df']['Marca'] == marca_selecionada]['PublicoAlvo'].iloc[0]
                opcoes_publico = [p.strip() for p in publicos_da_marca.split('.') if p.strip()]
                if not opcoes_publico: opcoes_publico = ["Público Geral (Baseado na Keyword)"]
                else: opcoes_publico.append("Público Geral (Baseado na Keyword)")
            except Exception:
                opcoes_publico = ["Público Geral (Baseado na Keyword)"]
                
            opcoes_publico.append("✍️ Digitar outro público (Personalizado)...")
            escolha_publico = st.selectbox("🎯 Para quem estamos escrevendo?", opcoes_publico)
            
            if escolha_publico == "✍️ Digitar outro público (Personalizado)...":
                publico_selecionado = st.text_input("Qual é o público-alvo?", placeholder="Ex: pais de alunos, estudantes do ensino médio...")
            else:
                publico_selecionado = escolha_publico
            # ----------------------------------------------
            # NOVOS INPUTS DO GERADOR
            # ----------------------------------------------
            # Inicializa a pauta se não existir
            if 'pauta_sugerida' not in st.session_state:
                st.session_state['pauta_sugerida'] = ""
            
            # ... dentro da coluna de inputs ...
            palavra_chave_input = st.text_area(
                "🔑 Palavra-chave ou Consulta/Query de Pesquisa", 
                value=st.session_state['pauta_sugerida'], # <--- Conecta aqui
                placeholder="Ex: metodologia bilíngue..."
            )
            
            conteudo_adicional_input = st.text_area(
                "📚 Conteúdo Adicional (Opcional)", 
                height=120,
                placeholder="Exemplos do que inserir aqui:\n- Links de referência: https://site.com/pesquisa-recente\n- Autores/Teorias: Cite a teoria de Vygotsky sobre o assunto.\n- Insumos próprios: 'Nossa escola parceira aumentou as matrículas em 20%...'\n- Restrições: Não fale sobre provas do MEC neste texto."
            )
            
            # ---> NOVO CAMPO: CONTEÚDO PROPRIETÁRIO <---
            conteudo_proprietario_input = st.text_area(
                "🔒 Conteúdo Proprietário Inegociável (Opcional)", 
                height=100,
                help="Frases exatas, citações ou parágrafos que a IA é OBRIGADA a incluir literalmente no texto gerado sem alterar nenhuma palavra.",
                placeholder="Ex: 'Segundo nosso diretor João, a educação transforma o amanhã.' (A IA vai colar este texto exato dentro do artigo)."
            )

            # ---> NOVO CAMPO: PROMPT LIVRE DO USUÁRIO <---
            instrucao_livre_input = st.text_area(
                "💬 Instruções Específicas / Prompt Livre (Estilo ChatGPT)", 
                height=120,
                help="Dite as regras do texto! Peça uma estrutura específica, perguntas exatas para os H2 ou um formato sob medida.",
                placeholder='Ex: "Preciso de um texto sobre o Vestibular da UERJ. Use uma estrutura de H2 respondendo: como funciona, o que cai, livros obrigatórios e a estrutura das provas."'
            )
            
            # O NOSSO NOVO INTERRUPTOR A/B
            st.markdown("<br>", unsafe_allow_html=True)
            modo_humanizado = st.toggle("✨ Ativar Escrita Empática / Mentoria (Beta)", value=False, help="Se ativado, a IA usa um prompt focado em fluidez humana e cadência vocal, reduzindo o tom corporativo.")
            
            # --- NOVO RECURSO: GHOSTWRITING ---
            modo_especialista = st.toggle("👔 Ativar Escrita de Especialista", value=False, help="A IA vai ler os artigos do especialista no Brandbook e escrever o texto usando o tom de voz, maneirismos e referências dele.")
            especialista_selecionado = None
            if modo_especialista:
                lista_autores = st.session_state['especialistas_df']['Especialista'].unique().tolist()
                especialista_selecionado = st.selectbox("Selecione o Autor/Especialista:", lista_autores)
            st.markdown("<br>", unsafe_allow_html=True)
    
            gerar_btn = st.button("🚀 Gerar Artigo em HTML", width="stretch", type="primary")
        
            st.markdown("---")
            
            cms_u, cms_usr, cms_p, cms_t = obter_credenciais_cms(marca_selecionada)
            WP_READY = bool(cms_u and cms_p)
    
            if not WP_READY:
                st.warning(f"🔌 Integração CMS inativa para a marca {marca_selecionada}. Faltam as credenciais no painel de Secrets.")
            else:
                # Faz um Ping real na API para ver se o Firewall está bloqueando
                with st.spinner(f"Verificando conexão com o Firewall do {cms_t.upper()}..."):
                    try:
                        import base64
                        token_teste = base64.b64encode(f"{cms_usr}:{cms_p.replace(' ', '').strip()}".encode('utf-8')).decode('utf-8')
                        # Máscara de Chrome para TODOS (WP e Drupal)
                        user_agent_ping = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
                        
                        headers_teste = {
                            'User-Agent': user_agent_ping, 
                            'Accept': 'application/json' if cms_t == 'wp' else 'application/vnd.api+json',
                            'Authorization': f'Basic {token_teste}',
                            'Connection': 'keep-alive'
                        }
                        
                        # Ping rápido puxando só 1 post (bem leve)
                        url_ping = f"{cms_u}?per_page=1" if cms_t == "wp" else f"{cms_u}?page[limit]=1"
                        res_ping = requests.get(url_ping, headers=headers_teste, timeout=5)
                        
                        if res_ping.status_code == 200:
                            st.success(f"🔌 Conectado e Autorizado no {cms_t.upper()} da marca: {marca_selecionada}")
                        elif res_ping.status_code in [403, 401]:
                            st.error(f"🛑 Credenciais OK, mas o Firewall (WAF) bloqueou a leitura da marca {marca_selecionada} (Erro {res_ping.status_code}). Solicite whitelist do User-Agent para a TI.")
                            WP_READY = False # Força o desativamento do botão de postagem direta mais abaixo
                        else:
                            st.warning(f"⚠️ API respondeu com Erro {res_ping.status_code}. O RAG Reverso pode falhar.")
                    except Exception:
                        st.error(f"🔌 O domínio da marca {marca_selecionada} não respondeu a tempo (Timeout).")
                        WP_READY = False

    # 2. DIRECIONANDO O CARREGAMENTO PARA A COLUNA 2 (DIREITA)
    if gerar_btn:
        if not TOKEN:
            st.error("⚠️ Erro: A chave OPENROUTER_KEY não foi encontrada nos Secrets.")
        elif not palavra_chave_input:
            st.warning("⚠️ Por favor, digite uma palavra-chave.")
        else:
            with col2: # <--- ISSO JOGA O SPINNER DE LOADING PARA A DIREITA
                with st.status("🤖 Processando Motor GEO v7.0...", expanded=True) as status:
                    try:
                        (
                            artigo_html, dicas_json, google_data, ia_data, entity_gap, 
                            score_originalidade, citabilidade, cluster, reverse_queries, 
                            citation_score, entity_coverage, geo_score, retrieval_simulation, 
                            hijacking_risk, ai_simulation, chunk_citability, answer_first, 
                            rag_chunks, evidence_density, information_gain, contexto_wp,
                            manual_voz_gemini
                        ) = executar_geracao_completa(
                            palavra_chave_input, marca_selecionada, publico_selecionado, 
                            conteudo_adicional_input, conteudo_proprietario_input, modo_humanizado, especialista_selecionado,
                            instrucao_livre_input # <--- ADICIONADO AQUI NO FINAL
                        )
                        
                        st.session_state['art_gerado'] = artigo_html
                        st.session_state['metas_geradas'] = dicas_json
                        st.session_state['google_ctx'] = google_data
                        st.session_state['ia_ctx'] = ia_data
                        st.session_state['entity_gap'] = entity_gap
                        st.session_state['score_originalidade'] = score_originalidade
                        st.session_state['citabilidade'] = citabilidade
                        st.session_state['cluster'] = cluster
                        st.session_state['reverse_queries'] = reverse_queries
                        st.session_state['citation_score'] = citation_score
                        st.session_state['entity_coverage'] = entity_coverage
                        st.session_state['geo_score'] = geo_score
                        st.session_state['retrieval_simulation'] = retrieval_simulation
                        st.session_state['hijacking_risk'] = hijacking_risk
                        st.session_state['ai_simulation'] = ai_simulation
                        st.session_state['chunk_citability'] = chunk_citability
                        st.session_state['answer_first'] = answer_first
                        st.session_state['rag_chunks'] = rag_chunks
                        st.session_state['evidence_density'] = evidence_density
                        st.session_state['information_gain'] = information_gain
                        st.session_state['contexto_wp'] = contexto_wp
                        st.session_state['manual_voz_gemini'] = manual_voz_gemini
                        st.session_state['marca_atual'] = marca_selecionada
                        st.session_state['keyword_atual'] = palavra_chave_input
                        status.update(label="✅ Artigo gerado com sucesso!", state="complete", expanded=False)
                    except Exception as e:
                        status.update(label="❌ Erro durante a geração", state="error")
                        st.error(f"Erro Crítico: {e}")

    if 'art_gerado' in st.session_state:
        with col2:
            st.success("✨ Tudo pronto! Seu artigo foi gerado e estruturado com sucesso.")
            
            # --- TÍTULO E MÉTRICA PRINCIPAL ---
            kpi_c1, kpi_c2 = st.columns(2)
            with kpi_c1:
                st.metric("🎯 Nota Geral de Estrutura (GEO)", st.session_state.get('citation_score', 'N/A'), help="Baseado em 5 critérios que o Google e as IAs mais valorizam hoje.")
            
            try:
                string_json_limpa = st.session_state['metas_geradas'].strip().removeprefix('```json').removesuffix('```').strip()
                meta_validada = MetadadosArtigo.model_validate_json(string_json_limpa)
                meta = meta_validada.model_dump()
                st.subheader(meta.get("title", "Artigo Gerado"))
            except Exception:
                meta = {"title": "Artigo Gerado (JSON Fallback)", "meta_description": "", "dicas_imagens": [], "schema_faq": {}}
                st.subheader("Artigo Gerado")

            st.markdown("<br>", unsafe_allow_html=True)

            # ==========================================
            # AS NOVAS SUB-ABAS DIDÁTICAS
            # ==========================================
            # 1. Declarando as abas (A de Leitura/Edição é a PRIMEIRA)
            tab_html, tab_dash, tab_seo, tab_ia = st.tabs([
                "👁️ Ler e Editar Artigo", 
                "📊 Dashboard Rápido", 
                "🧠 Raio-X Técnico de SEO", 
                "🤖 Como as IAs Enxergam"
            ])

            # --- SUB-ABA 1 (A PRIMEIRA): LER, EDITAR E COPIAR ---
            with tab_html:
                # O SELETOR DE MODO DE LEITURA VS EDIÇÃO
                modo_visualizacao = st.radio("O que você deseja fazer?", ["📖 Modo de Leitura", "✏️ Modo de Edição Manual"], horizontal=True, label_visibility="collapsed")
                
                if modo_visualizacao == "📖 Modo de Leitura":
                    
                    # O BOTÃO MÁGICO DE COPIAR TEXTO FORMATADO (JS CUSTOMIZADO)
                    componente_copiar = f"""
                    <div style="font-family: 'Inter', sans-serif;">
                        <button id="copy-btn" onclick="copyText()" style="background-color: #111827; color: white; border: none; padding: 12px 20px; border-radius: 8px; width: 100%; cursor: pointer; font-weight: 600; font-size: 15px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1); transition: all 0.3s;">
                            📋 Copiar Texto Formatado (Para colar no Docs/Word)
                        </button>
                        <div id="content-to-copy" style="position: absolute; left: -9999px;">
                            {st.session_state['art_gerado']}
                        </div>
                        <script>
                            function copyText() {{
                                var content = document.getElementById("content-to-copy");
                                var range = document.createRange();
                                range.selectNodeContents(content);
                                var selection = window.getSelection();
                                selection.removeAllRanges();
                                selection.addRange(range);
                                try {{
                                    document.execCommand("copy");
                                    var btn = document.getElementById("copy-btn");
                                    btn.innerHTML = "✅ Texto copiado com sucesso! Agora é só dar Ctrl+V no Docs.";
                                    btn.style.backgroundColor = "#10B981"; // Fica Verde
                                    setTimeout(function() {{
                                        btn.innerHTML = "📋 Copiar Texto Formatado (Para colar no Docs/Word)";
                                        btn.style.backgroundColor = "#111827"; // Volta pro Preto
                                    }}, 3000);
                                }} catch (err) {{
                                    console.error("Erro ao copiar: ", err);
                                }}
                                selection.removeAllRanges();
                            }}
                        </script>
                    </div>
                    """
                    # Renderiza o botão JS no Streamlit
                    st.components.v1.html(componente_copiar, height=65)
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.markdown("### 👁️ Pré-visualização do Blog")
                    
                    # A caixa envelopando o texto perfeitamente
                    html_preview = f"<div style='padding: 20px; border: 1px solid #E5E7EB; border-radius: 8px; background-color: #FFFFFF; color: #111827;'>{st.session_state['art_gerado']}</div>"
                    st.markdown(html_preview, unsafe_allow_html=True)
                    
                    st.markdown("---")
                    
                    # O código HTML fica escondidinho embaixo para quem quiser colar no painel do WordPress/Drupal
                    with st.expander("📋 Ver Código Fonte (HTML puro)"):
                        st.caption("Passe o mouse no canto superior direito da caixa preta abaixo e clique no ícone para copiar as tags HTML.")
                        st.code(st.session_state['art_gerado'], language="html")

                    # ==========================================
                    # NOVO RECURSO: CAIXA DE COMENTÁRIOS DA IA
                    # ==========================================
                    st.markdown("<br>", unsafe_allow_html=True)
                    st.markdown("### 🪄 Comentários para edição com IA")
                    st.caption("Escreva o que você deseja alterar no texto acima. A IA vai modificar apenas o trecho solicitado e manter o resto do artigo intacto.")
                    
                    # O campo de input em branco
                    instrucao_ajuste = st.text_area("Instruções:", placeholder="Ex: Substitua a palavra 'alunos' por 'estudantes' no segundo parágrafo... ou Remova o último tópico.", label_visibility="collapsed")
                    
                    # O botão só fica habilitado (clicável) se tiver algum texto na caixa
                    botao_habilitado = bool(instrucao_ajuste.strip())
                    
                    if st.button("✨ Refinar Texto com IA", type="secondary", disabled=not botao_habilitado, use_container_width=True):
                        with st.spinner("Cirurgia em andamento... A IA está reescrevendo apenas o trecho solicitado..."):
                            try:
                                novo_html = refinar_artigo_html(st.session_state['art_gerado'], instrucao_ajuste)
                                
                                # Guilhotina de segurança para limpar o markdown
                                novo_html = re.sub(r'^```html\n|```$', '', novo_html, flags=re.MULTILINE).strip()
                                if '<' in novo_html and '>' in novo_html:
                                    novo_html = novo_html[novo_html.find('<') : novo_html.rfind('>') + 1]
                                
                                # Atualiza o estado com o novo texto e recarrega a tela instantaneamente
                                st.session_state['art_gerado'] = novo_html
                                st.success("✅ Ajuste aplicado com sucesso!")
                                time.sleep(1)
                                st.rerun()
                            except Exception as e:
                                st.error(f"Erro ao tentar refinar o texto: {e}")

                else:
                    # MODO DE EDIÇÃO MANUAL DO CÓDIGO
                    st.markdown("### ✏️ Edição Manual de Texto/HTML")
                    st.caption("Faça as alterações que desejar no texto ou nas tags abaixo e clique em Salvar.")
                    
                    html_editado = st.text_area("Edite o conteúdo diretamente abaixo:", value=st.session_state['art_gerado'], height=450, label_visibility="collapsed")
                    
                    if st.button("💾 Salvar Edições Manuais", type="primary"):
                        st.session_state['art_gerado'] = html_editado
                        st.success("✅ Edições salvas com sucesso! Alterne para o 'Modo de Leitura' para ver o resultado.")
                        time.sleep(1.5)
                        st.rerun()

                # BOTÃO DE PUBLICAÇÃO DIRETA (Sempre visível no fundo da aba principal)
                st.markdown("<br>", unsafe_allow_html=True)
                cms_u, cms_usr, cms_p, cms_t = obter_credenciais_cms(st.session_state['marca_atual'])
                if cms_u and cms_usr and cms_p:
                    st.subheader(f"🌐 Publicação Direta ({cms_t.upper()})")
                    if st.button(f"📤 Enviar Rascunho para {cms_t.upper()} ({st.session_state['marca_atual']})", type="primary", width="stretch", key="btn_pub_principal"):
                        with st.spinner(f"Enviando via API para o {cms_t.upper()}..."):
                            if cms_t == "drupal":
                                res = publicar_drupal(meta.get("title", st.session_state['keyword_atual']), st.session_state['art_gerado'], meta, cms_u, cms_usr, cms_p)
                            elif cms_t == "webflow":
                                res = publicar_webflow(meta.get("title", st.session_state['keyword_atual']), st.session_state['art_gerado'], meta, cms_u, cms_usr, cms_p)
                            else:
                                res = publicar_wp(meta.get("title", st.session_state['keyword_atual']), st.session_state['art_gerado'], meta, cms_u, cms_usr, cms_p)
                            
                            if hasattr(res, 'status_code') and res.status_code in [200, 201]:
                                link_retorno = res.json().get('link') if hasattr(res, 'json') else "Rascunho criado!"
                                st.success(f"✅ Rascunho criado com sucesso! | {link_retorno}")
                            else:
                                erro_status = res.status_code if hasattr(res, 'status_code') else 'Desconhecido'
                                erro_texto = res.text if hasattr(res, 'text') else 'Sem detalhes'
                                st.error(f"❌ Falha ao enviar (Erro HTTP {erro_status}). Resposta do Servidor: {erro_texto}")

            # --- SUB-ABA 2: DASHBOARD RÁPIDO ---
            with tab_dash:
                st.info("**O que é esta aba?** Aqui estão as métricas essenciais para garantir que o seu texto será lido por humanos e ranqueado pelo Google.")
                
                with st.expander("🚀 Qualidade Global do Texto (GEO Score)", expanded=True):
                    st.markdown("Uma nota de 0 a 100 que resume se o texto está direto ao ponto, original e bem estruturado. **Acima de 80 é excelente.**")
                    st.json(st.session_state.get('geo_score', '{}'))
                    
                with st.expander("📑 O texto está fácil de ler? (Chunk Citability)", expanded=True):
                    st.markdown("IAs odeiam blocos gigantes de texto. Aqui medimos se o artigo tem **parágrafos curtos, listas e respostas diretas** logo no início (Answer-First).")
                    st.json(st.session_state.get('chunk_citability', '{}'))
                    st.json(st.session_state.get('answer_first', '{}'))

                with st.expander("🥇 O texto traz novidades? (Originalidade e Dados)", expanded=True):
                    st.markdown("O Google pune textos que só 'reciclam' o que já existe. Avaliamos se você trouxe **links de pesquisa, dados reais (Densidade de Evidências)** e palavras novas em relação aos concorrentes.")
                    st.json(st.session_state.get('evidence_density', '{}'))
                    st.json(st.session_state.get('information_gain', '{}'))
                    st.markdown(st.session_state.get('score_originalidade', '⚠️ Sem dados.'))

            # --- SUB-ABA 3: RAIO-X TÉCNICO DE SEO ---
            with tab_seo:
                st.info("**O que é esta aba?** Voltada para quem entende de SEO. Mostra se usamos o vocabulário certo e como amarrar este artigo com outros no seu blog.")
                
                # --- CAIXA DO GEMINI ADICIONADA AQUI ---
                with st.expander("🎭 Manual de Voz (Gerado pelo Gemini)", expanded=True):
                    st.markdown("Veja como o Gemini interpretou os PDFs da marca e quais regras ele passou para o Redator (Claude) imitar o estilo.")
                    st.info(st.session_state.get('manual_voz_gemini', '⚠️ Sem dados do Gemini.'))
                # ---------------------------------------    
                
                with st.expander("🧩 Uso de Jargões do Nicho (Entity Coverage)", expanded=True):
                    st.markdown("Avaliamos se o texto contém as 'Entidades' (termos técnicos e jargões) que provam para o Google que você é especialista no assunto, cobrindo buracos que os concorrentes deixaram (Entity Gap).")
                    st.json(st.session_state.get('entity_coverage', '{}'))
                    st.markdown(st.session_state.get('entity_gap', '⚠️ Sem dados.'))

                with st.expander("🗺️ Pautas Futuras (Content Cluster)", expanded=False):
                    st.markdown("Ideias de novos artigos que você pode escrever para linkar com este, criando uma 'Teia de Autoridade' no seu blog.")
                    st.markdown(st.session_state.get('cluster', '⚠️ Sem dados.'))
                    
                with st.expander("🔗 Linkagem Interna Automática", expanded=False):
                    st.markdown("O Motor vasculhou seu WordPress e obrigou a IA a linkar este artigo novo com posts antigos da sua marca para fortalecer seu SEO.")
                    st.markdown(st.session_state.get('contexto_wp', '⚠️ Sem dados.'))

            # --- SUB-ABA 4: COMO AS IAS ENXERGAM ---
            with tab_ia:
                st.info("**O que é esta aba?** Descubra se o ChatGPT ou o Perplexity usariam o seu texto como fonte oficial para responder a um usuário.")
                
                with st.expander("🔎 Chance de virar Fonte Oficial (Retrieval Simulation)", expanded=True):
                    st.markdown("A nossa simulação testa se o seu texto é neutro e confiável o suficiente para ser citado com link por uma Inteligência Artificial.")
                    st.json(st.session_state.get('retrieval_simulation', '{}'))
                    st.markdown(st.session_state.get('citabilidade', '⚠️ Sem dados.'))
                    
                with st.expander("⚠️ Risco de perder o leitor (Hijacking)", expanded=False):
                    st.markdown("Se o seu texto enrolar muito para explicar um conceito, uma IA concorrente pode 'roubar' sua citação simplesmente por ser mais didática. Avaliamos esse risco aqui.")
                    st.json(st.session_state.get('hijacking_risk', '{}'))
                    
                with st.expander("🤖 Teste Real: Como a resposta apareceria no ChatGPT", expanded=False):
                    st.markdown("Simulamos a tela do usuário final. Se ele perguntasse sobre esse tema para uma IA, é assim que a resposta seria gerada usando apenas o seu artigo como base.")
                    st.json(st.session_state.get('ai_simulation', '{}'))
                    
                with st.expander("🔄 O que as pessoas realmente perguntam? (Search Intent)", expanded=False):
                    st.markdown("Engenharia reversa: mapeamos as perguntas exatas que usuários leigos digitam no Google e as dúvidas profundas que a IA tenta resolver.")
                    st.json(st.session_state.get('reverse_queries', '{}'))
                            
# ==========================================
# 6. MONITOR DE GEO (GAMIFICAÇÃO E AUDITORIA)
# ==========================================
elif st.session_state['current_page'] == "Monitor de GEO":
    st.subheader("🔍 Monitor de Autoridade GEO")
    st.caption("Esta aba utiliza o **GPT-4o** para simular um algoritmo de busca, auditar seu texto e gerar insights estruturais.")
    
    conteudo_para_auditoria = st.session_state.get('art_gerado', '')
    keyword_para_auditoria = st.session_state.get('keyword_atual', '')
    marca_para_auditoria = st.session_state.get('marca_atual', 'a marca').replace('@', '')

    txt_auditoria = st.text_area("HTML do Artigo para Auditoria", height=300, value=conteudo_para_auditoria)
    kw_auditoria = st.text_input("Palavra-Chave Alvo", value=keyword_para_auditoria)

    if st.button("🔎 Analisar com GPT-4o e Gerar Insights"):
        if not txt_auditoria:
            st.warning("⚠️ Por favor, gere um artigo na aba 1 primeiro ou cole o HTML aqui.")
        else:
            with st.spinner("Executando Raio-X Matemático e Auditoria Semântica (GPT-4o)..."):
                
                # ==========================================
                # 1. EXECUTA A MATEMÁTICA NO TEXTO INSERIDO
                # ==========================================
                try:
                    math_chunk = avaliar_chunk_citability(txt_auditoria)
                    math_evidence = calcular_evidence_density(txt_auditoria)
                    math_answer = avaliar_answer_first(txt_auditoria)
                except Exception as e:
                    math_chunk, math_evidence, math_answer = {}, {}, {}
                    print(f"Erro nas métricas matemáticas do monitor: {e}")

                # ==========================================
                # 2. AUDITORIA SEMÂNTICA (GPT-4o)
                # ==========================================
                sys_audit = """Você é um Auditor Sênior de SEO e E-E-A-T do Google, além de Especialista em Engenharia de Prompt. Seu padrão é altíssimo, mas baseado em lógica estrutural e não em achismos.
                
                REGRAS CRÍTICAS DE AUDITORIA (VETOS ABSOLUTOS E PENALIZAÇÕES):
                1. AVALIAÇÃO DE CONCORRENTES VS. PARCEIROS: O texto é proibido de mencionar rivais comerciais da marca alvo. Tecnologias de apoio ou certificadoras da marca SÃO PARCEIROS.
                2. RASTREABILIDADE DE DADOS E LINKS: Se o texto apresentar dados DE MERCADO ou citar PESQUISAS e INSTITUIÇÕES (ex: British Council, USP), é obrigatório ter um link referencial real (<a href>). Sem link, REDUZA A NOTA. EXCEÇÃO ABSOLUTA: Dados institucionais da própria Marca Alvo (ex: % de fidelização, número de escolas) SÃO PROPRIETÁRIOS e não levam link.
                3. VALIDAÇÃO CONCEITUAL: Se o texto for puramente conceitual, NÃO requer links.
                4. TOM DA MARCA ALVO: A marca deve ser mencionada com tom de estudo de caso.
                5. IMAGENS IGNORADAS: IGNORE COMPLETAMENTE AS TAGS HTML DE IMAGEM (<img...>) NA SUA AVALIAÇÃO.
                6. LIBERDADE TEXTUAL (SEM NITPICKING): É expressamente proibido penalizar o texto, criar críticas ou reduzir pontos por causa de jargões corporativos, clichês ou expressões como "mundo globalizado", "cenário atual", "em resumo" ou "transcendeu". Foque exclusivamente na estrutura E-E-A-T e nos dados, e deixe o estilo literário livre.
                7. LINKAGEM DA MARCA (FOCO NO HTML): Toda vez que a Marca Alvo for mencionada no texto, ela deve obrigatoriamente conter um link (href) para sua URL oficial. Penalize se a marca for citada sem o link para o site.
                
                DIRETRIZ DE PONTUAÇÃO E FEEDBACK (A REGRA DOS 100 PONTOS):
                - Se você NÃO encontrar nenhuma quebra das regras acima (ou seja, se os arrays 'critica' e 'melhoria' estiverem vazios), O SCORE DEVE SER ESTRITAMENTE 100.
                - É EXPRESSAMENTE PROIBIDO subtrair pontos (ex: dar 90, 85) baseando-se em avaliações estéticas subjetivas ou vocabulário. Se tirou ponto, a justificativa TÊM que estar no array 'critica' baseada nas regras de 1 a 7.
                
                VOCÊ DEVE RETORNAR EXCLUSIVAMENTE UM OBJETO JSON COM A SEGUINTE ESTRUTURA:
                {
                  "score": "Um número inteiro de 0 a 100",
                  "veredito": "Resumo de autoridade.",
                  "critica": ["Ponto fraco 1", "Ponto fraco 2"],
                  "melhoria": ["Como arrumar 1"],
                  "sugestoes_dev": ["Insight para o prompt do Redator"]
                }"""
                
                usr_audit = f"""Palavra-chave: {kw_auditoria}
                Texto HTML: {txt_auditoria}
                Marca Alvo: {marca_para_auditoria}
                
                Audite e retorne APENAS o JSON."""
                
                try:
                    relatorio_bruto = chamar_llm(
                        sys_audit, 
                        usr_audit, 
                        model="openai/gpt-4o", 
                        temperature=0.1, 
                        response_format={"type": "json_object"}
                    )
                    
                    relatorio_limpo = relatorio_bruto.strip().removeprefix('```json').removesuffix('```').strip()
                    dados_audit = json.loads(relatorio_limpo)
                    
                    score = int(dados_audit.get("score", 0))
                    
                    st.markdown("---")
                    st.markdown("### 📊 Relatório de Performance GEO")
                    
                    # EXIBIÇÃO: Métricas Matemáticas vs Semânticas
                    col_math1, col_math2, col_math3 = st.columns(3)
                    with col_math1:
                        st.metric("📏 Chunk Citability (Estrutura)", f"{math_chunk.get('chunk_citability_score', 0)}/100", help="Mede a facilidade de IAs lerem o texto (listas e parágrafos curtos).")
                    with col_math2:
                        st.metric("⚡ Answer First", f"{math_answer.get('answer_first_score', 0)}/100", help="Verifica se a resposta direta está no topo do texto.")
                    with col_math3:
                        st.metric("🔗 Evidence Density", f"{math_evidence.get('evidence_density_score', 0)}/100", help="Volume de números e links no texto.")

                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    kpi1, kpi2 = st.columns([1, 3])
                    with kpi1:
                        cor_delta = "normal" if score >= 80 else "inverse"
                        st.metric("🎯 LLM Audit Score", f"{score}/100", delta=f"{score - 100} do ideal", delta_color=cor_delta, help="Nota dada pelo GPT-4o baseada nas regras de E-E-A-T.")
                    
                    with kpi2:
                        st.markdown("**Progresso E-E-A-T (Qualitativo):**")
                        st.progress(score / 100)
                        
                        if score == 100:
                            st.success(f"🏆 **Veredito de Autoridade:** {dados_audit.get('veredito')}")
                        elif score >= 80:
                            st.info(f"✅ **Veredito de Autoridade:** {dados_audit.get('veredito')}")
                        else:
                            st.warning(f"⚠️ **Veredito de Autoridade:** {dados_audit.get('veredito')}")

                    st.markdown("#### Análise Qualitativa (GPT-4o)")
                    col_critica, col_melhoria = st.columns(2)
                    
                    with col_critica:
                        with st.expander("🚨 Críticas Técnicas ao Texto", expanded=True):
                            criticas = dados_audit.get('critica', [])
                            if isinstance(criticas, list) and criticas:
                                for c in criticas:
                                    st.markdown(f"- {c}")
                            else:
                                st.markdown("✅ **Nenhuma crítica identificada. O texto passou ileso!**")
                                
                    with col_melhoria:
                        with st.expander("🛠️ Correções para este Artigo", expanded=True):
                            melhorias = dados_audit.get('melhoria', [])
                            if isinstance(melhorias, list) and melhorias:
                                for m in melhorias:
                                    st.markdown(f"- {m}")
                            else:
                                st.markdown("✅ **Sem sugestões de melhoria pendentes.**")

                    st.markdown("---")
                    st.markdown("### ⚙️ Engenharia de Prompt (Melhoria Contínua)")
                    with st.expander("💡 Sugestões de Novos Guardrails Estruturais", expanded=True):
                        sugestoes_dev = dados_audit.get('sugestoes_dev', [])
                        if isinstance(sugestoes_dev, list) and len(sugestoes_dev) > 0:
                            for s in sugestoes_dev:
                                st.info(f"🤖 **Insight para o Prompt:** {s}")
                            st.caption("Dica: Copie os insights acima que fizerem sentido e cole no prompt do Claude no código principal.")
                        else:
                            st.success("✨ **O prompt atual está performando de forma excelente para este nicho. Nenhuma sugestão gerada.**")

                except Exception as e:
                    st.error(f"Ocorreu um erro ao processar a auditoria visual. Detalhe técnico: {e}")
                    with st.expander("Ver resposta bruta da IA"):
                        st.write(relatorio_bruto if 'relatorio_bruto' in locals() else "Nenhuma resposta obtida.")

# ==========================================
# 7. ADAPTADOR DE PDF & REVISOR GEO (ABA 4)
# ==========================================
elif st.session_state['current_page'] == "Revisor de GEO":
    st.subheader("♻️ Adaptador & Revisor GEO")
    st.caption("Adapte um E-book/PDF proprietário para a voz de qualquer marca ou revise um artigo antigo do WordPress.")
    
    col_rev_1, col_rev_2 = st.columns([1, 2])
    
    with col_rev_1:
        marca_rev = st.selectbox("Selecione a Marca", st.session_state['brandbook_df']['Marca'].tolist(), key="marca_revisor")
        
        # --- Extração Dinâmica de Público ---
        try:
            publicos_da_marca_rev = st.session_state['brandbook_df'][st.session_state['brandbook_df']['Marca'] == marca_rev]['PublicoAlvo'].iloc[0]
            opcoes_publico_rev = [p.strip() for p in publicos_da_marca_rev.split('.') if p.strip()]
            if not opcoes_publico_rev: opcoes_publico_rev = ["Público Geral"]
            else: opcoes_publico_rev.append("Público Geral")
        except:
            opcoes_publico_rev = ["Público Geral"]
            
        opcoes_publico_rev.append("✍️ Digitar outro público (Personalizado)...")
        escolha_publico_rev = st.selectbox("🎯 Para quem o artigo será adaptado?", opcoes_publico_rev, key="pub_revisor")
        
        if escolha_publico_rev == "✍️ Digitar outro público (Personalizado)...":
            publico_rev = st.text_input("Qual é o público-alvo?", key="pub_manual_rev")
        else:
            publico_rev = escolha_publico_rev
            
        palavra_chave_rev = st.text_input("🔑 Palavra-chave foco", placeholder="Ex: eca digital")
    
    with col_rev_2:
        modo_input = st.radio("Origem do Conteúdo:", ["Puxar do WordPress", "Inserir HTML Manualmente", "Upload de Documentos (Base de Conhecimento)"], horizontal=True)
        conteudo_input = ""
        instrucoes_extras = "" # Inicializa a variável
        
        if modo_input == "Puxar do WordPress":
            url_r, user_r, pwd_r, type_r = obter_credenciais_cms(marca_rev)
            if url_r and pwd_r:
                if type_r == "drupal":
                    posts_cms = listar_posts_drupal(url_r, user_r, pwd_r)
                elif type_r == "webflow":
                    posts_cms = listar_posts_webflow(url_r, user_r, pwd_r)
                else:
                    posts_cms = listar_posts_wp(url_r, user_r, pwd_r)
                
                if posts_cms:
                    opcoes_posts = {f"{p['id']} - {p.get('title', {}).get('rendered', 'Sem Titulo')}": p.get('content', {}).get('rendered', '') for p in posts_cms}
                    post_selecionado = st.selectbox("Selecione o Artigo (Últimos 15 posts):", list(opcoes_posts.keys()))
                    conteudo_input = opcoes_posts[post_selecionado]
                    with st.expander("👁️ Ver HTML Original"):
                        st.code(conteudo_input[:1000] + "...\n(truncado)", language="html")
                else:
                    st.warning("⚠️ Nenhum post encontrado no CMS.")
            else:
                st.warning("🔌 Credenciais do CMS não configuradas.")
                
        elif modo_input == "Inserir HTML Manualmente":
            conteudo_input = st.text_area("Cole o HTML/Texto Original Aqui:", height=200)
            
        elif modo_input == "Upload de Documentos (Base de Conhecimento)":
            arquivos_upados = st.file_uploader("📄 Arraste seus arquivos (PDF, DOCX, TXT). Pode enviar mais de um!", type=['pdf', 'docx', 'txt'], accept_multiple_files=True)
            
            # ---> NOVO: AVISO DIDÁTICO PARA O USUÁRIO <---
            st.info("💡 **Dica de Uso:** Se você preencher o campo abaixo, a IA vai criar um artigo estruturado exatamente como você pedir. Se deixar em branco, ela criará automaticamente um texto 'Teaser/Spoiler' focado em captar leads para baixar o documento original.")
            
            instrucoes_extras = st.text_area(
                "✍️ Direcionamento (Prompt Complementar)", 
                height=150,
                placeholder="Ex: Eu tenho um documento em bullet points do time de produto. Preciso transformar esta descrição em um texto fácil de ler otimizado para o ENEM. Inclua os H2: O que é?, Como funciona? e Quais as vantagens? (quebradas em H3)."
            )
            
            if arquivos_upados:
                with st.spinner("Lendo base de documentos..."):
                    conteudo_input = extrair_texto_documentos(arquivos_upados)
                
                if "[Erro ao ler" in conteudo_input and len(conteudo_input.strip()) < 100:
                    st.error("Erro crítico ao ler os arquivos. Verifique se o pacote python-docx está no requirements.txt.")
                else:
                    st.success(f"✅ Documentos lidos com sucesso! ({len(conteudo_input)} caracteres extraídos).")
                    with st.expander("Ver Texto Bruto Extraído"):
                        st.text(conteudo_input[:2000] + "\n\n... (truncado)")

    if st.button("✨ Construir Artigo e Formatar (GEO)", type="primary", width="stretch"):
        if not TOKEN:
            st.error("⚠️ Chave OPENROUTER_KEY não encontrada.")
        elif not palavra_chave_rev or not conteudo_input:
            st.warning("⚠️ Preencha a palavra-chave e forneça o conteúdo base.")
        else:
            with st.spinner("Sintetizando documentos e construindo a estrutura... Isso pode levar alguns segundos."):
                try:
                    # Direcionamento do fluxo
                    if modo_input == "Upload de Documentos (Base de Conhecimento)":
                        resultado_processamento = executar_adaptacao_documentos(palavra_chave_rev, publico_rev, marca_rev, conteudo_input, instrucoes_extras)
                    else:
                        resultado_processamento = executar_revisao_geo_wp(palavra_chave_rev, publico_rev, marca_rev, conteudo_input)
                    
                    # Captura JSON da IA
                    match_json = re.search(r'\{.*\}', resultado_processamento.strip(), re.DOTALL)
                    json_limpo = match_json.group(0) if match_json else resultado_processamento.strip().removeprefix('```json').removesuffix('```').strip()
                    
                    try:
                        dados_processados = json.loads(json_limpo, strict=False)
                    except json.JSONDecodeError:
                        st.toast("⚠️ Corrigindo aspas duplas mal formatadas pela IA...", icon="🔧")
                        html_match = re.search(r'"html_novo"\s*:\s*"(.*?)"\s*\}?\s*$', json_limpo, re.DOTALL)
                        
                        html_resgatado = ""
                        if html_match:
                            html_resgatado = html_match.group(1).replace('\\"', '"').replace('\\n', '\n')
                            if html_resgatado.endswith('"}'): html_resgatado = html_resgatado[:-2]
                            elif html_resgatado.endswith('"'): html_resgatado = html_resgatado[:-1]
                        
                        dados_processados = {
                            "diagnostico": "JSON recuperado via fallback de Regex.",
                            "melhorias_aplicadas": ["Correção forçada de formatação"],
                            "html_novo": html_resgatado if html_resgatado else "<p>Erro crítico de formatação da IA. Tente gerar novamente.</p>"
                        }
                    
                    st.success("Adaptação concluída com sucesso!")
                    
                    col_resultado_1, col_resultado_2 = st.columns(2)
                    
                    with col_resultado_1:
                        st.markdown("### 📋 Diagnóstico da Adaptação")
                        st.info(f"**O que foi feito:**\n{dados_processados.get('diagnostico', 'N/A')}")
                        
                        st.markdown("### 🛠️ Melhorias Aplicadas")
                        for m in dados_processados.get('melhorias_aplicadas', []):
                            st.markdown(f"- ✅ {m}")
                    
                    with col_resultado_2:
                        st.markdown("### 🚀 Novo Código HTML")
                        st.code(dados_processados.get('html_novo', ''), language="html")
                        
                    st.markdown("---")
                    st.markdown("### 👁️ Pré-visualização do Artigo")
                    
                    # ---> NOVO: BOTÃO DE COPIAR EMBUTIDO <---
                    html_para_copiar = dados_processados.get('html_novo', '')
                    
                    componente_copiar_rev = f"""
                    <div style="font-family: 'Inter', sans-serif;">
                        <button id="copy-btn-rev" onclick="copyTextRev()" style="background-color: #111827; color: white; border: none; padding: 12px 20px; border-radius: 8px; width: 100%; cursor: pointer; font-weight: 600; font-size: 15px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1); transition: all 0.3s;">
                            📋 Copiar Texto Formatado (Para colar no Docs/Word)
                        </button>
                        <div id="content-to-copy-rev" style="position: absolute; left: -9999px;">
                            {html_para_copiar}
                        </div>
                        <script>
                            function copyTextRev() {{
                                var content = document.getElementById("content-to-copy-rev");
                                var range = document.createRange();
                                range.selectNodeContents(content);
                                var selection = window.getSelection();
                                selection.removeAllRanges();
                                selection.addRange(range);
                                try {{
                                    document.execCommand("copy");
                                    var btn = document.getElementById("copy-btn-rev");
                                    btn.innerHTML = "✅ Texto copiado com sucesso! Agora é só dar Ctrl+V no Docs.";
                                    btn.style.backgroundColor = "#10B981"; // Fica Verde
                                    setTimeout(function() {{
                                        btn.innerHTML = "📋 Copiar Texto Formatado (Para colar no Docs/Word)";
                                        btn.style.backgroundColor = "#111827"; // Volta pro Preto
                                    }}, 3000);
                                }} catch (err) {{
                                    console.error("Erro ao copiar: ", err);
                                }}
                                selection.removeAllRanges();
                            }}
                        </script>
                    </div>
                    """
                    
                    st.components.v1.html(componente_copiar_rev, height=65)
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    html_preview_rev = f"<div style='padding: 20px; border: 1px solid #E5E7EB; border-radius: 8px; background-color: #FFFFFF; color: #111827;'>{html_para_copiar}</div>"
                    st.markdown(html_preview_rev, unsafe_allow_html=True)
                    
                except Exception as e:
                    st.error(f"Erro ao processar: {e}")

# ==========================================
# 8. AUDITOR DE ARTIGOS (NOVA ABA 5)
# ==========================================
elif st.session_state['current_page'] == "Auditor de Artigos":
    st.subheader("📊 Auditor de Visibilidade GEO")
    st.caption("Verifique se um artigo publicado está ranqueando no Google ou sendo recomendado espontaneamente pelas IAs.")

    col_aud_1, col_aud_2 = st.columns([1, 2])
    
    with col_aud_1:
        marca_auditor = st.selectbox("Marca Analisada", st.session_state['brandbook_df']['Marca'].tolist(), key="marca_auditor_tab5")
        palavra_chave_auditor = st.text_input("🔑 Palavra-chave Alvo", placeholder="Ex: metodologia bilíngue")
        
        modo_url = st.radio("Origem do Artigo:", ["Puxar do CMS", "Inserir Manualmente"], horizontal=True)
        
        url_auditor = ""
        if modo_url == "Puxar do CMS":
            cms_u_aud, cms_usr_aud, cms_p_aud, cms_t_aud = obter_credenciais_cms(marca_auditor)
            
            if cms_u_aud and cms_p_aud:
                with st.spinner(f"Buscando os últimos artigos publicados no {cms_t_aud.upper()}..."):
                    if cms_t_aud == "drupal":
                        posts_aud = listar_posts_drupal(cms_u_aud, cms_usr_aud, cms_p_aud)
                    elif cms_t_aud == "webflow":
                        posts_aud = listar_posts_webflow(cms_u_aud, cms_usr_aud, cms_p_aud)
                    else:
                        posts_aud = listar_posts_wp(cms_u_aud, cms_usr_aud, cms_p_aud)
                    
                if posts_aud:
                    opcoes_url = {}
                    for p in posts_aud:
                        tit = p.get('title', {}).get('rendered', 'Sem Título')
                        link_post = p.get('link', '')
                        opcoes_url[f"{p.get('id')} - {tit}"] = link_post
                        
                    post_sel = st.selectbox("🔗 Selecione o Artigo Publicado:", list(opcoes_url.keys()))
                    url_auditor = opcoes_url[post_sel]
                    
                    if url_auditor:
                        st.caption(f"**URL Selecionada:** `{url_auditor}`")
                    else:
                        st.warning("⚠️ O CMS não retornou a URL para este post. Pode ser um rascunho.")
                        url_auditor = st.text_input("🔗 Digite a URL manualmente", key="url_manual_fallback")
                else:
                    st.warning("⚠️ Nenhum post encontrado ou bloqueio de Firewall. Tente digitar manualmente.")
                    url_auditor = st.text_input("🔗 URL do Artigo", key="url_manual_fallback_2")
            else:
                st.warning("🔌 Credenciais da marca não configuradas. Digite a URL manualmente.")
                url_auditor = st.text_input("🔗 URL do Artigo", key="url_manual_no_creds")
        else:
            url_auditor = st.text_input("🔗 URL do Artigo Publicado", placeholder="Ex: https://www.saseducacao.com.br/artigo-teste", key="url_manual_direto")
        
    with col_aud_2:
        st.info("""
        💡 **Como o Auditor funciona?**
        1. Ele faz uma **engenharia reversa** baseada nas possíveis buscas (*Reverse Query Engine - Search Intent*) a partir da palavra-chave.
        2. **Pesquisa as possíveis buscas** em tempo real no Google e nos LLMs.
        3. Varre os resultados procurando a **URL** do seu artigo (ou o nome da marca nas IAs).
        4. **Avalia se o conteúdo precisa ir para a Revisão GEO** com base nos resultados.
        """)
        
    if st.button("🚀 Iniciar Auditoria de Visibilidade (Google e IA)", type="primary", width="stretch"):
        if not TOKEN or not SERPAPI_KEY:
            st.error("⚠️ As chaves de API estão faltando nos Secrets.")
        elif not palavra_chave_auditor:
            st.warning("⚠️ Preencha a palavra-chave principal para iniciarmos o Intent Map.")
        elif not url_auditor:
            st.warning("⚠️ Forneça a URL do artigo para podermos rastreá-la nas buscas.")
        else:
            with st.status("🕵️‍♂️ Iniciando Auditoria GEO Avançada...", expanded=True) as status_aud:
                
                # Passo 1: Engenharia Reversa (Search Intent)
                st.write("1️⃣ Analisando Intenção de Busca e gerando variações profundas...")
                rev_queries_str = gerar_reverse_queries(palavra_chave_auditor)
                
                try:
                    rev_data = json.loads(rev_queries_str)
                    
                    # Pega as 4 primeiras perguntas que os usuários reais fazem
                    uq = rev_data.get("user_questions", [])[:4]
                    
                    # Pega 1 pergunta do LLM só para garantir o contexto semântico
                    lrq = rev_data.get("llm_reasoning_questions", [])[:1]
                    
                    buscas_extras = uq + lrq
                    
                    # Remove duplicatas e limpa vazios
                    buscas_extras = [q for q in list(set(buscas_extras)) if q]
                    
                    buscas_alvo = [palavra_chave_auditor] + buscas_extras
                except Exception as e:
                    buscas_alvo = [palavra_chave_auditor]
                    rev_data = {"Erro": f"Não foi possível expandir as buscas. Detalhe: {e}"}
                
                st.write(f"🎯 Buscas que serão rastreadas: *{', '.join(buscas_alvo)}*")
                
                # Passo 2: Pesquisar as buscas no Google e LLMs em paralelo
                st.write("2️⃣ Rastreadores vasculhando o Google e consultando LLMs...")
                resultados_google_agregados = ""
                resultados_ia_agregados = ""
                
                def auditar_termo(termo):
                    return buscar_contexto_google(termo), buscar_baseline_llm(termo)
                
                # Executa múltiplas buscas simultâneas para ser rápido
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    futures = {executor.submit(auditar_termo, q): q for q in buscas_alvo}
                    for future in concurrent.futures.as_completed(futures):
                        q = futures[future]
                        try:
                            g_res, ia_res = future.result(timeout=45)
                            resultados_google_agregados += f"\n\n--- Busca: '{q}' ---\n{g_res}"
                            resultados_ia_agregados += f"\n\n--- Busca: '{q}' ---\n{ia_res}"
                        except Exception as e:
                            st.write(f"⚠️ Timeout/Erro ao buscar o termo '{q}': {e}")

                # Passo 3: Varrer os resultados procurando a URL
                st.write("3️⃣ Cruzando dados e procurando a URL fornecida...")
                
                url_limpa = url_auditor.lower().replace("https://", "").replace("http://", "").replace("www.", "").strip()
                # Tira a barra final se existir para não falhar no match do Google
                if url_limpa.endswith('/'): url_limpa = url_limpa[:-1] 
                
                marca_limpa = marca_auditor.lower().replace(" ", "")
                
                google_ranqueia_url = url_limpa in resultados_google_agregados.lower()
                ia_menciona_marca = marca_limpa in resultados_ia_agregados.lower().replace(" ", "")
                
                status_aud.update(label="✅ Auditoria Concluída!", state="complete", expanded=False)

            # ==================================
            # Passo 4: Avaliação e Revisor GEO
            # ==================================
            st.markdown("---")
            st.subheader("🎯 Veredito da Auditoria")
            
            c_res1, c_res2 = st.columns(2)
            with c_res1:
                st.markdown("### 🌐 Google Search (Múltiplas Buscas)")
                if google_ranqueia_url:
                    st.success(f"✅ **SUCESSO EXTREMO!** A URL do seu artigo foi encontrada no Top 3 orgânico ou Featured Snippets para as buscas testadas.")
                else:
                    st.error(f"❌ **NÃO RANQUEIA.** O Google não está mostrando sua URL no Top 3 para o Intent Map gerado.")
            
            with c_res2:
                st.markdown("### 🤖 Consenso de IA (Share of Voice)")
                if ia_menciona_marca:
                    st.success(f"✅ **AUTORIDADE RECONHECIDA!** As IAs estão citando a marca **{marca_auditor}** espontaneamente nestes tópicos.")
                else:
                    st.error(f"❌ **PONTO CEGO DE IA.** O consenso das inteligências artificiais não cita a sua marca como autoridade nessas buscas.")
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            # Lógica de Encaminhamento
            if not google_ranqueia_url or not ia_menciona_marca:
                st.error("🚨 **ALERTA DE DESEMPENHO:** Este conteúdo está perdendo tráfego e visibilidade para os motores gerativos.")
                st.markdown(f"> **Ação Recomendada:** O seu texto não está atendendo aos critérios de E-E-A-T ou Answer-First esperados pelas IAs. Vá para a aba **📝 Revisor GEO WordPress**, selecione este artigo e peça para o Motor reescrevê-lo aplicando as métricas matemáticas corretas.")
            else:
                st.success("🏆 **BLINDAGEM TOTAL:** Seu artigo é uma fortaleza GEO. Está dominando o Google Clássico e as recomendações de IA. Nenhuma ação de revisão é necessária!")

            # Expansores de Dados
            st.markdown("---")
            with st.expander("🔄 Mapa de Intenção de Busca (Search Intent Gap) Utilizado", expanded=False):
                st.caption("As perguntas abaixo revelam o que os usuários e as IAs realmente querem saber sobre esse tema. Se seu texto não responde a isso, precisa de revisão.")
                st.json(rev_data)

            with st.expander("🕵️‍♂️ Auditoria Bruta: O que ranqueia hoje (Google & IA)?"):
                st.markdown("**O que os Robôs do Google leram no Top 3 durante as buscas:**")
                st.code(resultados_google_agregados if resultados_google_agregados else "Sem dados.", language="markdown")
                
                st.markdown("**Como as IAs responderam às perguntas hoje:**")
                st.code(resultados_ia_agregados if resultados_ia_agregados else "Sem dados.", language="markdown")

st.markdown("""
<div style="text-align: center; color: #6b7280; font-size: 13px; margin-top: 60px; padding-top: 20px; border-top: 1px solid #e5e7eb; line-height: 1.8;">
    ⚙️ <strong>Feito para simplificar o complexo.</strong> Criação otimizada para humanos e novos motores de busca.<br>
    ⚙️ <strong>Stack:</strong> Python | Streamlit | Pydantic &nbsp;&nbsp;&nbsp;&nbsp; 🧠 <strong>LLMs:</strong> GPT-4o | Claude 4 Sonnet | Gemini 2.5 Pro &nbsp;&nbsp;&nbsp;&nbsp; 🔌 <strong>APIs:</strong> Serper.dev | Jina AI | Unsplash
</div>
""", unsafe_allow_html=True)
